# =============================================================================
# DASHBOARD CARTERA IPS — Clínica Mar Caribe / COLSALUD
# Ejecutar con: streamlit run app.py
# Requiere: pip install streamlit pandas plotly openpyxl xlrd requests python-docx
# =============================================================================

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import requests
import warnings
warnings.filterwarnings("ignore")

st.set_page_config(
    page_title="Tablero de Cartera — IPS",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

PALETTE = ["#1B3A6B","#2E86AB","#F4A261","#2D9B6F","#E63946",
           "#A8DADC","#457B9D","#E9C46A","#264653","#E76F51"]
COLORS  = {"primary":"#1B3A6B","secondary":"#2E86AB","success":"#2D9B6F",
           "warning":"#F4A261","danger":"#E63946","orange":"#E76F51"}

st.markdown("""
<style>
.stApp{background:#F0F2F6}
.main-header{background:linear-gradient(135deg,#1B3A6B,#2E86AB);padding:1.2rem 2rem;
  border-radius:12px;margin-bottom:1rem}
.main-header h1{color:white;margin:0;font-size:1.7rem;font-weight:700}
.main-header p{color:rgba(255,255,255,.85);margin:.2rem 0 0;font-size:.9rem}
.kpi-card{background:white;border-radius:12px;padding:1rem 1.2rem;
  box-shadow:0 2px 8px rgba(0,0,0,.08);border-left:5px solid #2E86AB;margin-bottom:.8rem}
.kpi-card.danger{border-left-color:#E63946}
.kpi-card.warning{border-left-color:#F4A261}
.kpi-card.success{border-left-color:#2D9B6F}
.kpi-card.primary{border-left-color:#1B3A6B}
.kpi-card.purple{border-left-color:#7B2FBE}
.kpi-card.orange{border-left-color:#E76F51}
.kpi-label{font-size:.75rem;font-weight:600;color:#6C757D;text-transform:uppercase;
  letter-spacing:.05em;margin-bottom:.2rem}
.kpi-value{font-size:1.4rem;font-weight:700;color:#1B3A6B;line-height:1.2}
.kpi-sub{font-size:.76rem;color:#6C757D;margin-top:.15rem}
.section-title{font-size:1.05rem;font-weight:700;color:#1B3A6B;
  border-bottom:2px solid #2E86AB;padding-bottom:.3rem;margin:1.2rem 0 .8rem}
.insight-box{background:#EEF4FB;border-left:4px solid #2E86AB;border-radius:8px;
  padding:.8rem 1rem;margin-bottom:.6rem;font-size:.87rem;color:#1B3A6B}
.insight-box.danger{background:#FEF0F0;border-color:#E63946;color:#7B1D1D}
.insight-box.warning{background:#FFF8EC;border-color:#F4A261;color:#7A4A00}
.insight-box.success{background:#EDFDF6;border-color:#2D9B6F;color:#1A5B3A}
.analisis-box{background:white;border-radius:10px;padding:1rem 1.2rem;
  box-shadow:0 2px 6px rgba(0,0,0,.07);margin-bottom:.8rem;border-left:4px solid #2E86AB;
  font-size:.88rem;color:#333;line-height:1.7}
section[data-testid="stSidebar"]{background-color:#1B3A6B}
section[data-testid="stSidebar"] *{color:white!important}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def fmt(v):
    try: return f"${int(round(v)):,}".replace(",",".")
    except: return "$0"
def fmt_pct(v):
    try: return f"{v:.1f}%"
    except: return "0.0%"
def kpi_card(label, value, sub="", kind="default"):
    cls={"danger":"danger","warning":"warning","success":"success",
         "primary":"primary","purple":"purple","orange":"orange"}.get(kind,"")
    st.markdown(f"""<div class="kpi-card {cls}">
        <div class="kpi-label">{label}</div>
        <div class="kpi-value">{value}</div>
        <div class="kpi-sub">{sub}</div></div>""", unsafe_allow_html=True)
def section(t):
    st.markdown(f'<div class="section-title">{t}</div>', unsafe_allow_html=True)
def insight(txt, kind="info"):
    icons={"danger":"🔴","warning":"🟡","success":"🟢","info":"🔵"}
    st.markdown(f'<div class="insight-box {kind if kind!="info" else ""}">{icons.get(kind,"🔵")} {txt}</div>',
                unsafe_allow_html=True)
def analisis_box(txt):
    st.markdown(f'<div class="analisis-box">{txt}</div>', unsafe_allow_html=True)
def fig_layout(fig, title="", height=370):
    fig.update_layout(
        title=dict(text=title,font=dict(size=13,color="#1B3A6B"),x=0),
        height=height,paper_bgcolor="white",plot_bgcolor="white",
        font=dict(family="Arial",size=11,color="#333"),
        legend=dict(orientation="h",yanchor="bottom",y=-0.35,xanchor="center",x=0.5),
        margin=dict(l=10,r=10,t=38,b=10),
    )
    fig.update_xaxes(showgrid=False,linecolor="#DEE2E6")
    fig.update_yaxes(gridcolor="#F0F0F0",linecolor="#DEE2E6")
    return fig

_kc = {"n": 0}
def plot(fig, **kwargs):
    _kc["n"] += 1
    st.plotly_chart(fig, key=f"chart_{_kc['n']}", **kwargs)

# ── Columnas opcionales ───────────────────────────────────────────────────────
COLS_OPT = {
    "glosa_valor":    "Vr,Glosa Reportada",
    "glosa_saldo":    "Saldo de Glosa según cartera",
    "glosa_estado":   "Estado De Glosa",
    "glosa_dias":     "Dias de Glosa",
    "glosa_intervalo":"Intervalo Dias Glosa",
    "juridico_estado":"Estado juridico",
    "juridico_proceso":"Proceso Juridico",
    "abogado":        "ABOGADO",
    "vencida":        "Factura vencida",
    "por_vencer":     "Por vencer",
    "dias_mora":      "Dias Mora",
    "observacion":    "OBSERVACION",
    "supersalud":     "ACUERDOS SUPERSALUD Y OTROS",
    "nit":            "Nuevo Nit",
    "fecha_venc":     "Fecha de Vencimiento",
    "fecha_fac":      "Fecha fac",
    "fecha_rad":      "Fecha rad",
    "recepcion":      "Fecha_Recepcion",
    "objetivo_dias":  "Objetivo Menos de x Dias",
    "vr_factura":     "Vr factura",
    "devuelta":       "Devuelta",
    "dev_juridica":   "Facturas devueltas por juridica",
    "motivo_dev":     "Motivo de devolucion",
    "vr_acuerdo":     "VR ACUERDO",
    "pgp":            "PGP",
    "año_fac":        "Año Facturacion",
    "año_rad":        "Año de radicacion",
    "copago":         "Copago",
    "vr_aceptado":    "Vr.Aceptado",
    "vr_pagado":      "Vr.Pagado",
}
COLS_REQ = ["Razon Social","Tipo De Empresa","Saldo Actual","Estado de cartera","Intervalo-Actual"]

def tiene(df, key):
    col = COLS_OPT.get(key)
    return col is not None and col in df.columns

def to_num(s):
    return pd.to_numeric(s.astype(str).str.replace(r"[^\d\.\-]","",regex=True),
                         errors="coerce").fillna(0)

# ══════════════════════════════════════════════════════════════════════════════
# CARGA
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data(show_spinner=False)
def cargar_cartera(file):
    try: df = pd.read_excel(file, dtype=str)
    except Exception as e: return None, str(e)
    df.columns = [c.strip() for c in df.columns]
    falt = [c for c in COLS_REQ if c not in df.columns]
    if falt: return None, f"Columnas faltantes: {falt}"
    df["Saldo Actual"] = to_num(df["Saldo Actual"])
    for k in ["glosa_valor","glosa_saldo","dias_mora","glosa_dias","vr_factura","vr_acuerdo","vr_aceptado","vr_pagado","copago"]:
        col = COLS_OPT.get(k)
        if col and col in df.columns: df[col] = to_num(df[col])
    for k in ["vencida","por_vencer"]:
        col = COLS_OPT.get(k)
        if col and col in df.columns: df[f"_num_{k}"] = to_num(df[col])
    for k in ["fecha_venc","fecha_fac","fecha_rad"]:
        col = COLS_OPT.get(k)
        if col and col in df.columns:
            df[col] = pd.to_datetime(df[col], errors="coerce")
    col_ff = COLS_OPT.get("fecha_fac"); col_fr = COLS_OPT.get("fecha_rad")
    if col_ff and col_fr and col_ff in df.columns and col_fr in df.columns:
        df["_dias_rad"] = (df[col_fr] - df[col_ff]).dt.days.clip(lower=0)
    col_obj = COLS_OPT.get("objetivo_dias")
    if col_obj and col_obj in df.columns and "_dias_rad" in df.columns:
        obj_num = to_num(df[col_obj])
        df["_dentro_objetivo"] = df["_dias_rad"] <= obj_num
    return df, None

@st.cache_data(show_spinner=False)
def cargar_pagos(file):
    try: dp = pd.read_excel(file, dtype=str)
    except Exception as e: return None, str(e)
    dp.columns = [c.strip() for c in dp.columns]
    rm = {}
    for c in dp.columns:
        cl = c.lower().strip()
        if "nit" in cl and "razon" not in cl: rm[c]="Nit"
        elif "razon" in cl or "social" in cl:  rm[c]="Razon Social"
        elif "fecha" in cl:                    rm[c]="Fecha de pago"
        elif any(x in cl for x in ["valor","pago","monto"]): rm[c]="Valor pagado"
        elif "estado" in cl:                   rm[c]="Estado"
    dp = dp.rename(columns=rm)
    falt = [c for c in ["Nit","Razon Social","Fecha de pago","Valor pagado","Estado"] if c not in dp.columns]
    if falt: return None, f"Columnas no encontradas: {falt}"
    dp["Valor pagado"]  = to_num(dp["Valor pagado"])
    dp["Fecha de pago"] = pd.to_datetime(dp["Fecha de pago"], errors="coerce")
    dp["Nit"]           = dp["Nit"].astype(str).str.strip()
    return dp, None

# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
def build_sidebar(df, tab_activa="resumen"):
    st.sidebar.markdown("## 🔎 Filtros")
    def ms(label, col):
        if col in df.columns:
            opts = sorted(df[col].dropna().astype(str).unique().tolist())
            return st.sidebar.multiselect(label, opts, default=[])
        return []
    f_razon  = ms("Razón Social",      "Razon Social")
    f_tipo   = ms("Tipo de Empresa",   "Tipo De Empresa")
    f_estado = ms("Estado de Cartera", "Estado de cartera")
    f_interv = ms("Intervalo-Actual",  "Intervalo-Actual")
    f_año_fac= ms("Año Facturación",   "Año Facturacion")
    f_año_rad= ms("Año Radicación",    "Año de radicacion")
    f_obs = ""
    if "OBSERVACION" in df.columns:
        f_obs = st.sidebar.text_input("🔍 Buscar Observación","")
    f_glosa_estado = f_glosa_interv = []
    if tab_activa == "glosas":
        st.sidebar.markdown("---")
        st.sidebar.markdown("**🔍 Filtros de Glosas**")
        if tiene(df,"glosa_estado"):    f_glosa_estado = ms("Estado Glosa",        COLS_OPT["glosa_estado"])
        if tiene(df,"glosa_intervalo"): f_glosa_interv = ms("Intervalo Días Glosa", COLS_OPT["glosa_intervalo"])
    f_juridico = f_abogado = f_proceso = []
    if tab_activa == "juridico":
        st.sidebar.markdown("---")
        st.sidebar.markdown("**⚖️ Filtros Jurídico**")
        if tiene(df,"juridico_estado"):  f_juridico = ms("Estado Jurídico",  COLS_OPT["juridico_estado"])
        if tiene(df,"abogado"):          f_abogado  = ms("Abogado",          COLS_OPT["abogado"])
        if tiene(df,"juridico_proceso"): f_proceso  = ms("Proceso Jurídico", COLS_OPT["juridico_proceso"])
    st.sidebar.markdown("---")
    st.sidebar.metric("Total registros", f"{len(df):,}")
    st.sidebar.metric("Clientes únicos", f"{df['Razon Social'].nunique():,}")
    return dict(razon=f_razon,tipo=f_tipo,estado=f_estado,interv=f_interv,
                año_fac=f_año_fac,año_rad=f_año_rad,obs=f_obs,
                glosa_estado=f_glosa_estado,glosa_interv=f_glosa_interv,
                juridico=f_juridico,abogado=f_abogado,proceso=f_proceso)

def aplicar_filtros(df, f):
    d = df.copy()
    if f["razon"]:   d = d[d["Razon Social"].isin(f["razon"])]
    if f["tipo"]:    d = d[d["Tipo De Empresa"].isin(f["tipo"])]
    if f["estado"]:  d = d[d["Estado de cartera"].isin(f["estado"])]
    if f["interv"]:  d = d[d["Intervalo-Actual"].isin(f["interv"])]
    if f["año_fac"] and "Año Facturacion" in d.columns: d = d[d["Año Facturacion"].isin(f["año_fac"])]
    if f["año_rad"] and "Año de radicacion" in d.columns: d = d[d["Año de radicacion"].isin(f["año_rad"])]
    if f["obs"] and "OBSERVACION" in d.columns:
        d = d[d["OBSERVACION"].astype(str).str.contains(f["obs"],case=False,na=False)]
    if f["glosa_estado"] and tiene(d,"glosa_estado"):
        d = d[d[COLS_OPT["glosa_estado"]].isin(f["glosa_estado"])]
    if f["glosa_interv"] and tiene(d,"glosa_intervalo"):
        d = d[d[COLS_OPT["glosa_intervalo"]].isin(f["glosa_interv"])]
    if f["juridico"] and tiene(d,"juridico_estado"):
        d = d[d[COLS_OPT["juridico_estado"]].isin(f["juridico"])]
    if f["abogado"] and tiene(d,"abogado"):
        d = d[d[COLS_OPT["abogado"]].isin(f["abogado"])]
    if f["proceso"] and tiene(d,"juridico_proceso"):
        d = d[d[COLS_OPT["juridico_proceso"]].isin(f["proceso"])]
    return d

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR WORD PARA NOTEBOOKLM
# ══════════════════════════════════════════════════════════════════════════════
def _generar_word_interno(df, dp):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None, "Librería python-docx no instalada. Ejecuta: pip install python-docx"

    import io
    doc   = Document()
    total = df["Saldo Actual"].sum()
    n_cli = df["Razon Social"].nunique()
    fecha = pd.Timestamp.now().strftime("%d de %B de %Y")

    def h1(txt):
        p = doc.add_heading(txt, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x1B,0x3A,0x6B); return p
    def h2(txt):
        p = doc.add_heading(txt, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x2E,0x86,0xAB); return p
    def parr(txt):
        p = doc.add_paragraph(); r = p.add_run(txt)
        r.font.size = Pt(11); return p
    def sep(): doc.add_paragraph("─"*80)
    def tabla_df(data, max_rows=20):
        if isinstance(data, dict): data = pd.DataFrame(data)
        data = data.head(max_rows)
        t = doc.add_table(rows=1, cols=len(data.columns))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(data.columns):
            hdr[i].text = str(col)
            run = hdr[i].paragraphs[0].runs
            if run: run[0].bold = True
        for _, row in data.iterrows():
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        doc.add_paragraph()

    # Portada
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("INFORME DE CARTERA — IPS")
    r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(0x1B,0x3A,0x6B)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run(f"Clínica Mar Caribe / COLSALUD\nCorte: {fecha}\nAnálisis integral para Comité de Cartera y Junta Directiva").font.size=Pt(12)
    doc.add_paragraph(); sep()

    # Métricas base
    cobro_df    = df[df["Estado de cartera"].astype(str).str.strip()=="Para cobro"]
    total_cobro = cobro_df["Saldo Actual"].sum()
    pct_cobro   = (total_cobro/total*100) if total else 0
    jur_df      = df[df["Estado de cartera"].astype(str).str.strip()=="En proceso juridico"]
    total_jur   = jur_df["Saldo Actual"].sum()
    liq_df      = df[df["Estado de cartera"].astype(str).str.lower().str.contains("liquidaci",na=False)]
    total_liq   = liq_df["Saldo Actual"].sum()
    pct_liq     = (total_liq/total*100) if total else 0
    total_glosa = df[COLS_OPT["glosa_saldo"]].sum() if tiene(df,"glosa_saldo") else 0
    cartera_neta= total - total_glosa
    por_radicar = df[df["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()
    top10       = df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10)
    conc_top10  = (top10.sum()/total*100) if total else 0
    tv          = df["_num_vencida"].sum() if "_num_vencida" in df.columns else 0
    v_crit      = df[df["Intervalo-Actual"].astype(str).str.strip().isin(
                    ["2 años","3 años","4 años","5 años","> 5 AÑOS","211-360"])]["Saldo Actual"].sum()

    # 1. Resumen ejecutivo
    h1("1. RESUMEN EJECUTIVO")
    parr(f"La cartera total al {fecha} asciende a {fmt(total)}, distribuida en {n_cli} clientes. "
         f"La cartera neta (descontando glosas {fmt(total_glosa)}) es {fmt(cartera_neta)}.")
    tabla_df({"Indicador":["Cartera Bruta","Cartera Neta","Glosas (saldo)","En Cobro","% Cobro",
                           "Vencida","Jurídica","Por Radicar","En Liquidación","Conc. Top 10"],
              "Valor":[fmt(total),fmt(cartera_neta),fmt(total_glosa),fmt(total_cobro),fmt_pct(pct_cobro),
                       fmt(tv),fmt(total_jur),fmt(por_radicar),fmt(total_liq),fmt_pct(conc_top10)]})

    # 2. Composición por estado
    h1("2. COMPOSICIÓN POR ESTADO")
    ec=df.groupby("Estado de cartera")["Saldo Actual"].sum().reset_index()
    ec["% Part"]=(ec["Saldo Actual"]/total*100).round(1)
    ec["Saldo"]=ec["Saldo Actual"].apply(fmt)
    ec=ec.sort_values("Saldo Actual",ascending=False)
    tabla_df(ec[["Estado de cartera","Saldo","% Part"]].rename(columns={"Estado de cartera":"Estado"}))
    parr(f"'Para cobro' concentra {fmt_pct(pct_cobro)} ({fmt(total_cobro)}). "
         f"Clientes en liquidación: {fmt(total_liq)} ({fmt_pct(pct_liq)}).")

    # 3. Aging
    h1("3. ENVEJECIMIENTO (AGING)")
    ORDEN=["Por radicar","1-30","31-60","61-90","91-150","151-210","211-360",
           "2 años","3 años","4 años","5 años","> 5 AÑOS"]
    ag=df.groupby("Intervalo-Actual")["Saldo Actual"].sum().reset_index()
    ag["_o"]=ag["Intervalo-Actual"].apply(
        lambda v:next((i for i,o in enumerate(ORDEN) if str(v).strip().upper()==o.upper()),99))
    ag=ag.sort_values("_o").drop(columns=["_o"])
    ag["% Part"]=(ag["Saldo Actual"]/total*100).round(1)
    ag["Saldo"]=ag["Saldo Actual"].apply(fmt)
    tabla_df(ag[["Intervalo-Actual","Saldo","% Part"]].rename(columns={"Intervalo-Actual":"Intervalo"}))
    parr(f"Cartera mayor a 211 días: {fmt(v_crit)} ({fmt_pct(v_crit/total*100 if total else 0)}). "
         f"Alta probabilidad de irrecuperabilidad. Evaluar provisión NIIF.")

    # 4. Top 10 deudores
    h1("4. TOP 10 MAYORES DEUDORES")
    t10=df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10).reset_index()
    t10["% Part"]=(t10["Saldo Actual"]/total*100).round(1)
    t10["Saldo"]=t10["Saldo Actual"].apply(fmt)
    t10.index=range(1,len(t10)+1)
    tabla_df(t10[["Razon Social","Saldo","% Part"]].rename(columns={"Razon Social":"Cliente"}))
    parr(f"Top 10 concentra {fmt_pct(conc_top10)} ({fmt(top10.sum())}). Mayor deudor: {top10.idxmax()} con {fmt(top10.max())}.")

    # 5. Tipo de empresa
    h1("5. ANÁLISIS POR TIPO DE EMPRESA")
    te=df.groupby("Tipo De Empresa")["Saldo Actual"].sum().reset_index()
    te["% Part"]=(te["Saldo Actual"]/total*100).round(1)
    te["Saldo"]=te["Saldo Actual"].apply(fmt)
    tabla_df(te[["Tipo De Empresa","Saldo","% Part"]].rename(columns={"Tipo De Empresa":"Tipo"}))

    # 6. Glosas
    h1("6. ANÁLISIS DE GLOSAS")
    cg=COLS_OPT.get("glosa_valor"); cgs=COLS_OPT.get("glosa_saldo"); col_vf=COLS_OPT.get("vr_factura")
    if cg and cg in df.columns:
        dfg=df[df[cg]>0].copy()
        tg=dfg[cgs].sum() if (cgs and cgs in dfg.columns) else dfg[cg].sum()
        tv_g=dfg[cg].sum()
        idx_g=(tv_g/df[col_vf].sum()*100) if (col_vf and col_vf in df.columns and df[col_vf].sum()>0) else 0
        tabla_df({"Indicador":["Saldo Glosas s/Cartera","Vr Glosa Reportada","% sobre Cartera","Índice de Glosas","Registros","Clientes"],
                  "Valor":[fmt(tg),fmt(tv_g),fmt_pct(tg/total*100 if total else 0),f"{idx_g:.2f}%",f"{len(dfg):,}",str(dfg["Razon Social"].nunique())]})
        if col_vf and col_vf in df.columns:
            h2("6.1 Índice de Glosas por Cliente (Top 10)")
            ig=df[df[col_vf]>0].groupby("Razon Social").agg(Facturado=(col_vf,"sum"),Glosado=(cg,"sum")).reset_index()
            ig["Índice %"]=(ig["Glosado"]/ig["Facturado"]*100).round(2)
            ig=ig.sort_values("Índice %",ascending=False).head(10)
            ig["Facturado"]=ig["Facturado"].apply(fmt); ig["Glosado"]=ig["Glosado"].apply(fmt)
            tabla_df(ig[["Razon Social","Facturado","Glosado","Índice %"]].rename(columns={"Razon Social":"Cliente"}))
        parr(f"Las glosas representan {fmt_pct(tg/total*100 if total else 0)} de la cartera. "
             f"Índice de glosas: {idx_g:.2f}%. {'ALERTA: supera el 10% crítico.' if idx_g>10 else 'Dentro del rango aceptable (<10%).'}")

    # 7. Vencimientos
    h1("7. VENCIMIENTOS Y ALERTAS")
    tv_v=df["_num_vencida"].sum() if "_num_vencida" in df.columns else 0
    tpv_v=df["_num_por_vencer"].sum() if "_num_por_vencer" in df.columns else 0
    tabla_df({"Indicador":["Factura Vencida","Por Vencer","Exposición Total"],
              "Valor":[fmt(tv_v),fmt(tpv_v),fmt(tv_v+tpv_v)],
              "% Total":[fmt_pct(tv_v/total*100 if total else 0),
                         fmt_pct(tpv_v/total*100 if total else 0),
                         fmt_pct((tv_v+tpv_v)/total*100 if total else 0)]})
    h2("7.1 Alertas por Intervalo")
    for interv,nivel in {"1-30":"CRÍTICO","31-60":"ALTO","61-90":"MEDIO","91-150":"MEDIO"}.items():
        di=df[df["Intervalo-Actual"].astype(str).str.strip()==interv]
        if not di.empty:
            si=di["Saldo Actual"].sum()
            top=di.groupby("Razon Social")["Saldo Actual"].sum().nlargest(3)
            cli=", ".join([f"{c} ({fmt(v)})" for c,v in top.items()])
            parr(f"• {nivel} — {interv}: {fmt(si)} ({fmt_pct(si/total*100 if total else 0)}). Clientes: {cli}.")

    # 8. Radicación
    if "_dias_rad" in df.columns:
        h1("8. EFICIENCIA DE RADICACIÓN")
        dr=df[df["_dias_rad"].notna()&(df["_dias_rad"]>=0)]
        prom_r=dr["_dias_rad"].mean(); med_r=dr["_dias_rad"].median(); max_r=dr["_dias_rad"].max()
        tabla_df({"Indicador":["Días Promedio","Mediana","Máximo","Facturas Analizadas"],
                  "Valor":[f"{prom_r:.0f} días",f"{med_r:.0f} días",f"{max_r:.0f} días",f"{len(dr):,}"]})
        parr(f"Promedio {prom_r:.0f} días. {'Supera el umbral de 30 días recomendado.' if prom_r>30 else 'Dentro del rango aceptable.'}")
        h2("8.1 Top 10 con Mayor Demora")
        top_dem=dr.groupby("Razon Social")["_dias_rad"].mean().nlargest(10).reset_index()
        top_dem.columns=["Cliente","Días Promedio"]; top_dem["Días Promedio"]=top_dem["Días Promedio"].round(0).astype(int)
        tabla_df(top_dem)

    # 9. Devueltas
    col_dev=COLS_OPT.get("devuelta"); col_mot=COLS_OPT.get("motivo_dev")
    if col_dev and col_dev in df.columns:
        h1("9. FACTURAS DEVUELTAS")
        dfd=df[df[col_dev].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])]
        if not dfd.empty:
            t_d=dfd["Saldo Actual"].sum()
            parr(f"{len(dfd):,} facturas devueltas por {fmt(t_d)} ({fmt_pct(t_d/total*100 if total else 0)}), "
                 f"afectando a {dfd['Razon Social'].nunique()} clientes.")
            top_dv=dfd.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10).reset_index()
            top_dv.columns=["Cliente","Saldo"]; top_dv["Saldo"]=top_dv["Saldo"].apply(fmt)
            tabla_df(top_dv)
            if col_mot and col_mot in dfd.columns:
                h2("9.1 Por Motivo")
                mot=dfd.groupby(col_mot)["Saldo Actual"].sum().reset_index()
                mot.columns=["Motivo","Saldo"]; mot["Saldo"]=mot["Saldo"].apply(fmt)
                tabla_df(mot.head(10))

    # 10. Jurídico
    col_je=COLS_OPT.get("juridico_estado"); col_ab=COLS_OPT.get("abogado")
    if col_je and col_je in df.columns:
        h1("10. ANÁLISIS JURÍDICO")
        dfj=df[~df[col_je].astype(str).str.strip().str.lower().isin(["","nan","no aplica"])]
        if not dfj.empty:
            tj=dfj["Saldo Actual"].sum()
            parr(f"Cartera en proceso jurídico: {fmt(tj)} ({fmt_pct(tj/total*100 if total else 0)}), {len(dfj):,} registros.")
            je=dfj.groupby(col_je)["Saldo Actual"].sum().reset_index()
            je.columns=["Estado Jurídico","Saldo"]; je["Saldo"]=je["Saldo"].apply(fmt)
            tabla_df(je)
            if col_ab and col_ab in dfj.columns:
                h2("10.1 Por Abogado")
                ja=dfj.groupby(col_ab)["Saldo Actual"].sum().nlargest(10).reset_index()
                ja.columns=["Abogado","Saldo"]; ja["Saldo"]=ja["Saldo"].apply(fmt)
                tabla_df(ja)

    # 11. Pagos
    if dp is not None:
        h1("11. PAGOS HISTÓRICOS")
        tp=dp["Valor pagado"].sum()
        reg=dp[dp["Estado"].astype(str).str.lower()=="registrado"]["Valor pagado"].sum()
        sin=dp[dp["Estado"].astype(str).str.lower()=="sin registrar"]["Valor pagado"].sum()
        tabla_df({"Indicador":["Total Recibido","Registrados","Recaudo por Aplicar","% Aplicado","Clientes"],
                  "Valor":[fmt(tp),fmt(reg),fmt(sin),fmt_pct(reg/tp*100 if tp else 0),str(dp["Razon Social"].nunique())]})
        h2("11.1 Top 10 Recaudo por Aplicar")
        ts=dp[dp["Estado"].astype(str).str.lower()=="sin registrar"].groupby("Razon Social")["Valor pagado"].sum().nlargest(10).reset_index()
        ts.columns=["Cliente","Sin Aplicar"]; ts["Sin Aplicar"]=ts["Sin Aplicar"].apply(fmt)
        tabla_df(ts)

    # 12. Cosecha
    col_af=COLS_OPT.get("año_fac")
    if col_af and col_af in df.columns:
        h1("12. COSECHA DE FACTURACIÓN")
        cos=df.groupby(col_af).agg(Saldo=("Saldo Actual","sum"),Facturas=("Saldo Actual","count")).reset_index()
        cos.columns=["Año","Saldo","Facturas"]
        cos["% Part"]=(cos["Saldo"]/total*100).round(1)
        cos["Saldo Fmt"]=cos["Saldo"].apply(fmt)
        cos=cos.sort_values("Año")
        tabla_df(cos[["Año","Saldo Fmt","Facturas","% Part"]].rename(columns={"Saldo Fmt":"Saldo"}))

    # 13. Hallazgos y recomendaciones
    h1("13. HALLAZGOS Y RECOMENDACIONES")
    h2("13.1 Hallazgos")
    top10p=(top10.sum()/total*100) if total else 0
    hs=[]
    if top10p>60: hs.append(f"ALTA CONCENTRACIÓN: Top 10 = {top10p:.1f}% de {n_cli} clientes. Riesgo sistémico de liquidez.")
    if pct_cobro>30: hs.append(f"COBRO CRÍTICO: {pct_cobro:.1f}% ({fmt(total_cobro)}) en estado 'Para cobro'. Acción inmediata.")
    if total_liq>0: hs.append(f"LIQUIDACIÓN: {liq_df['Razon Social'].nunique()} clientes = {fmt(total_liq)} ({fmt_pct(pct_liq)}). Alta probabilidad de pérdida.")
    if tiene(df,"glosa_saldo"):
        pg=(total_glosa/total*100) if total else 0
        if pg>5: hs.append(f"GLOSAS ELEVADAS: {fmt_pct(pg)} de la cartera. Supera umbral del 5%.")
    if "_dias_rad" in df.columns:
        pr=df[df["_dias_rad"].notna()]["_dias_rad"].mean()
        if pr>30: hs.append(f"DEMORA RADICACIÓN: {pr:.0f} días promedio. Supera el umbral de 30 días.")
    if v_crit>0: hs.append(f"CARTERA ANTIGUA: {fmt(v_crit)} ({fmt_pct(v_crit/total*100 if total else 0)}) con +211 días. Evaluar provisión NIIF.")
    for i,h in enumerate(hs,1): parr(f"{i}. {h}")

    h2("13.2 Recomendaciones Estratégicas")
    recs=[
        ("Priorizar gestión de cobro", f"Activar cobro sobre {len(cobro_df):,} registros en 'Para cobro' ({fmt(total_cobro)}). Asignar gestores por EPS."),
        ("Gestión Top deudores", f"Negociar acuerdos escalonados. Mayor deudor: {top10.idxmax()}."),
        ("Reducir tiempo de radicación", "Meta: menos de 15 días desde fecha de factura. Implementar radicación electrónica."),
        ("Gestión de glosas", "Cronograma quincenal con auditoría médica. Priorizar por valor y fecha."),
        ("Revisión jurídica mensual", "Informe de avance de cada apoderado. Evaluar acuerdos ante SuperSalud."),
        ("Provisión cartera antigua", f"Evaluar provisión NIIF 9 sobre {fmt(v_crit)} con antigüedad +211 días."),
        ("Acuerdos SuperSalud/ADRES", "Gestionar conciliación extrajudicial y compensaciones con ADRES."),
        ("Comité mensual de cartera", "KPIs mensuales: cobro, aging, concentración, jurídico, glosas, pagos."),
    ]
    for titulo,texto in recs:
        p=doc.add_paragraph(style="List Bullet")
        r1=p.add_run(f"{titulo}: "); r1.bold=True; r1.font.size=Pt(11)
        r2=p.add_run(texto); r2.font.size=Pt(11)

    # Pie
    doc.add_paragraph(); sep()
    pie=doc.add_paragraph(); pie.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pie.add_run(f"Generado el {fecha} | Tablero de Cartera IPS | Compatible con NotebookLM").font.size=Pt(9)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf, None

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
def seccion_resumen(df, dp):
    total=df["Saldo Actual"].sum(); n_cli=df["Razon Social"].nunique()
    ec_df=df.groupby("Estado de cartera")["Saldo Actual"].sum().reset_index()
    ec_df["Pct"]=(ec_df["Saldo Actual"]/total*100).round(1)
    ec_df=ec_df.sort_values("Saldo Actual",ascending=False)
    cobro_df=df[df["Estado de cartera"].astype(str).str.strip()=="Para cobro"]
    total_cobro=cobro_df["Saldo Actual"].sum(); pct_cobro=(total_cobro/total*100) if total else 0
    juridico_df=df[df["Estado de cartera"].astype(str).str.strip()=="En proceso juridico"]
    total_jur=juridico_df["Saldo Actual"].sum(); pct_jur=(total_jur/total*100) if total else 0
    liq_df=df[df["Estado de cartera"].astype(str).str.lower().str.contains("liquidaci",na=False)]
    total_liq=liq_df["Saldo Actual"].sum(); pct_liq=(total_liq/total*100) if total else 0
    total_glosa=df[COLS_OPT["glosa_saldo"]].sum() if tiene(df,"glosa_saldo") else 0
    cartera_neta=total-total_glosa
    por_radicar=df[df["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()
    pct_radicar=(por_radicar/total*100) if total else 0
    top10=df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10)
    conc_top10=(top10.sum()/total*100) if total else 0
    total_vencida=df["_num_vencida"].sum() if "_num_vencida" in df.columns else 0

    section("📝 Análisis General de Cartera")
    estados_txt="".join([f"• <b>{r['Estado de cartera']}</b>: {fmt(r['Saldo Actual'])} ({r['Pct']:.1f}%)<br>" for _,r in ec_df.iterrows()])
    analisis_box(f"<b>Corte:</b> {pd.Timestamp.now().strftime('%d de %B de %Y')}<br><br>"
                 f"Cartera total: <b>{fmt(total)}</b> | {n_cli} clientes | Cartera neta: <b>{fmt(cartera_neta)}</b><br><br>"
                 f"<b>Composición:</b><br>{estados_txt}<br>"
                 f"'Para cobro': <b>{fmt_pct(pct_cobro)}</b> ({fmt(total_cobro)}). "
                 f"Liquidación: <b>{fmt(total_liq)}</b> ({fmt_pct(pct_liq)}). "
                 f"Jurídico: <b>{fmt(total_jur)}</b> ({fmt_pct(pct_jur)}). "
                 f"Por radicar: <b>{fmt(por_radicar)}</b> ({fmt_pct(pct_radicar)}).")

    section("📊 Composición por Estado")
    bar_colors=[COLORS["danger"] if str(x).strip()=="Para cobro"
                else COLORS["orange"] if "liquidaci" in str(x).lower()
                else PALETTE[i%len(PALETTE)] for i,x in enumerate(ec_df["Estado de cartera"])]

    # Tabla siempre visible
    ec_tabla = ec_df.copy()
    ec_tabla["Saldo"] = ec_tabla["Saldo Actual"].apply(fmt)
    st.dataframe(
        ec_tabla[["Estado de cartera","Saldo","Pct"]].rename(
            columns={"Estado de cartera":"Estado","Pct":"% Part."}),
        use_container_width=True, hide_index=True)

    # Gráficos dentro de expander para evitar error DOM
    with st.expander("📊 Ver gráficos de composición"):
        cg1,cg2=st.columns([1.4,1])
        with cg1:
            fig=go.Figure(go.Bar(x=ec_df["Estado de cartera"],y=ec_df["Saldo Actual"],
                                 text=ec_df["Saldo Actual"].apply(fmt),textposition="outside",marker_color=bar_colors))
            plot(fig_layout(fig,"Saldo por Estado",320),use_container_width=True)
        with cg2:
            fig2=px.pie(ec_df,names="Estado de cartera",values="Saldo Actual",color_discrete_sequence=PALETTE,hole=0.42)
            fig2.update_traces(textposition="inside",textinfo="percent+label")
            plot(fig_layout(fig2,"Participación",320),use_container_width=True)

    section("💰 Indicadores Financieros Clave")
    c1,c2,c3,c4=st.columns(4)
    with c1: kpi_card("Cartera Bruta Total",fmt(total),f"{n_cli} clientes","primary")
    with c2: kpi_card("Cartera Neta",fmt(cartera_neta),"Bruta menos glosas","success")
    with c3: kpi_card("Glosas (saldo)",fmt(total_glosa),fmt_pct(total_glosa/total*100 if total else 0),"warning")
    with c4: kpi_card("Cartera en Cobro",fmt(total_cobro),fmt_pct(pct_cobro),"danger")
    c5,c6,c7,c8=st.columns(4)
    with c5: kpi_card("Cartera Vencida",fmt(total_vencida),"col. Factura vencida","warning")
    with c6: kpi_card("Cartera Jurídica",fmt(total_jur),fmt_pct(pct_jur),"danger")
    with c7: kpi_card("Por Radicar",fmt(por_radicar),fmt_pct(pct_radicar),"warning")
    with c8: kpi_card("Clientes en Liquidación",fmt(total_liq),f"{liq_df['Razon Social'].nunique()} clientes","orange")

    if dp is not None:
        section("💳 Resumen de Pagos")
        tp=dp["Valor pagado"].sum()
        reg=dp[dp["Estado"].astype(str).str.strip().str.lower()=="registrado"]
        sinreg=dp[dp["Estado"].astype(str).str.strip().str.lower()=="sin registrar"]
        pct_ap=(reg["Valor pagado"].sum()/tp*100) if tp else 0
        cp1,cp2,cp3,cp4=st.columns(4)
        with cp1: kpi_card("Total Pagos",fmt(tp),f"{len(dp):,} registros","purple")
        with cp2: kpi_card("Registrados",fmt(reg["Valor pagado"].sum()),f"{len(reg):,}","success")
        with cp3: kpi_card("Recaudo por Aplicar",fmt(sinreg["Valor pagado"].sum()),f"{len(sinreg):,}","danger")
        with cp4: kpi_card("% Aplicado",fmt_pct(pct_ap),"del total","success" if pct_ap>70 else "warning")
        with st.expander("📊 Ver gráficos de pagos"):
            if dp["Fecha de pago"].notna().any():
                dp2=dp.copy(); dp2["Año"]=dp2["Fecha de pago"].dt.year
                dp2["Mes_Num"]=dp2["Fecha de pago"].dt.month; dp2["Mes_Nom"]=dp2["Fecha de pago"].dt.strftime("%b")
                def etiq(a):
                    try: return "Acum. 2012-2023" if int(a)<=2023 else str(a)
                    except: return str(a)
                dp2["Grupo"]=dp2["Año"].apply(etiq)
                g1=dp2.groupby(["Grupo","Estado"])["Valor pagado"].sum().reset_index()
                og=[g for g in ["Acum. 2012-2023","2024","2025","2026"] if g in g1["Grupo"].unique()]
                g1["_o"]=g1["Grupo"].apply(lambda x:og.index(x) if x in og else 99)
                g1=g1.sort_values("_o").drop(columns=["_o"]); g1["Fmt"]=g1["Valor pagado"].apply(fmt)
                fp1=px.bar(g1,x="Grupo",y="Valor pagado",color="Estado",barmode="group",
                           color_discrete_map={"Registrado":"#2D9B6F","Sin registrar":"#E63946"},
                           text="Fmt",category_orders={"Grupo":og})
                fp1.update_traces(textposition="outside")
                plot(fig_layout(fp1,"Pagos por período",340),use_container_width=True)
                d26=dp2[dp2["Año"].isin([2025,2026])].copy()
                if not d26.empty:
                    MO=["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
                    g2=d26.groupby(["Año","Mes_Num","Mes_Nom"])["Valor pagado"].sum().reset_index()
                    g2=g2.sort_values(["Año","Mes_Num"]); g2["Año"]=g2["Año"].astype(str); g2["Fmt"]=g2["Valor pagado"].apply(fmt)
                    fp2=px.bar(g2,x="Mes_Nom",y="Valor pagado",color="Año",barmode="group",
                               color_discrete_map={"2025":"#2E86AB","2026":"#E63946"},
                               text="Fmt",category_orders={"Mes_Nom":MO})
                    fp2.update_traces(textposition="outside")
                    plot(fig_layout(fp2,"Comparativo 2025 vs 2026",340),use_container_width=True)

    section("⚠️ Alertas del Período")
    alertas=[]
    top1=df.groupby("Razon Social")["Saldo Actual"].sum().idxmax()
    val1=df.groupby("Razon Social")["Saldo Actual"].sum().max()
    alertas.append(("danger",f"Mayor deudor: <b>{top1}</b> — {fmt(val1)} ({fmt_pct(val1/total*100)})"))
    if total_cobro>0: alertas.append(("danger",f"<b>Para cobro</b>: {fmt(total_cobro)} ({fmt_pct(pct_cobro)}). Acción inmediata."))
    if total_liq>0:   alertas.append(("danger",f"<b>{liq_df['Razon Social'].nunique()} clientes en liquidación</b>: {fmt(total_liq)} ({fmt_pct(pct_liq)})."))
    if total_jur>0:   alertas.append(("warning",f"Cartera jurídica: {fmt(total_jur)} ({fmt_pct(pct_jur)})."))
    if por_radicar>0: alertas.append(("warning",f"Por radicar: {fmt(por_radicar)} ({fmt_pct(pct_radicar)}). Flujo de caja detenido."))
    if total_glosa>0: alertas.append(("warning",f"Saldo glosas: {fmt(total_glosa)} ({fmt_pct(total_glosa/total*100 if total else 0)})."))
    if dp is not None:
        sr=dp[dp["Estado"].astype(str).str.lower()=="sin registrar"]["Valor pagado"].sum()
        if sr>0: alertas.append(("warning",f"Recaudo por aplicar: {fmt(sr)}. Conciliación urgente."))
    for kind,txt in alertas:
        st.markdown(f'<div class="insight-box {kind}">{"🔴" if kind=="danger" else "🟡"} {txt}</div>',unsafe_allow_html=True)

    section("🏆 Top 7 Mayores Deudores")
    t7=df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(7).reset_index()
    t7["% Part."]=(t7["Saldo Actual"]/total*100).round(1); t7["Saldo"]=t7["Saldo Actual"].apply(fmt)
    t7.index=range(1,len(t7)+1)
    st.dataframe(t7[["Razon Social","Saldo","% Part."]].rename(columns={"Razon Social":"Cliente"}),use_container_width=True)
    with st.expander("📊 Ver gráfico Top 7"):
        f3=px.bar(t7.sort_values("Saldo Actual"),x="Saldo Actual",y="Razon Social",orientation="h",
                  color="Saldo Actual",color_continuous_scale=["#A8DADC","#2E86AB","#1B3A6B"],text="Saldo")
        f3.update_traces(textposition="outside"); f3.update_coloraxes(showscale=False)
        plot(fig_layout(f3,"",300),use_container_width=True)

    return total, pct_cobro, (total_vencida/total*100 if total else 0), conc_top10

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — KPIs + CONSULTOR
# ══════════════════════════════════════════════════════════════════════════════
def seccion_kpis(df, dp, total, pct_cobro, pct_vencida, conc_top10):
    section("📈 KPIs Financieros")
    prom=df.groupby("Razon Social")["Saldo Actual"].sum().mean(); n=df["Razon Social"].nunique()
    c1,c2,c3,c4,c5=st.columns(5)
    with c1: kpi_card("% Cartera Vencida",fmt_pct(pct_vencida),"Sobre total","warning")
    with c2: kpi_card("% En Cobro",fmt_pct(pct_cobro),"Estado Para cobro","danger")
    with c3: kpi_card("Concentración Top 10",fmt_pct(conc_top10),"vs total","primary")
    with c4: kpi_card("Promedio x Cliente",fmt(prom),f"{n} clientes","default")
    with c5:
        if tiene(df,"supersalud"):
            ss=to_num(df[COLS_OPT["supersalud"]]).sum()
            kpi_card("Acuerdos SuperSalud",fmt(ss),"Total acordado","success")
        else:
            kpi_card("Clientes Únicos",f"{n:,}","en cartera","default")
    st.markdown("---")
    section("🤖 Consultor Inteligente de Cartera")
    st.info("💡 Ejemplos de preguntas:\n\n"
            "• ¿Cuál es la factura de mayor saldo de SANITAS?\n\n"
            "• ¿Qué clientes tienen más de 180 días de mora?\n\n"
            "• Dame sugerencias sobre COOSALUD\n\n"
            "• ¿Cuánto suman los pagos sin registrar?")
    pregunta=st.text_area("✍️ Escribe tu pregunta:",height=90,placeholder="Ej: ¿Cuál es el cliente con mayor cartera vencida?")
    if st.button("🔍 Consultar",type="primary") and pregunta.strip():
        with st.spinner("Analizando..."):
            resp=_consultar_ia(pregunta,df,dp)
        st.divider(); st.markdown("**🤖 Respuesta:**")
        st.info(resp.replace("<b>","").replace("</b>","").replace("<br>","\n\n"))

def _construir_contexto(df, dp):
    total=df["Saldo Actual"].sum(); ctx=[]
    ctx.append(f"CARTERA TOTAL: {fmt(total)} | {len(df):,} registros | {df['Razon Social'].nunique()} clientes")
    ctx.append("ESTADOS: "+" | ".join([f"{k}: {fmt(v)}" for k,v in df.groupby("Estado de cartera")["Saldo Actual"].sum().nlargest(8).items()]))
    ctx.append("TOP 10: "+" | ".join([f"{k}: {fmt(v)}" for k,v in df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10).items()]))
    ctx.append("AGING: "+" | ".join([f"{k}: {fmt(v)}" for k,v in df.groupby("Intervalo-Actual")["Saldo Actual"].sum().items()]))
    if tiene(df,"glosa_saldo"): ctx.append(f"GLOSAS: {fmt(df[COLS_OPT['glosa_saldo']].sum())}")
    if "_num_vencida" in df.columns: ctx.append(f"VENCIDO: {fmt(df['_num_vencida'].sum())}")
    if dp is not None:
        tp=dp["Valor pagado"].sum()
        reg=dp[dp["Estado"].astype(str).str.lower()=="registrado"]["Valor pagado"].sum()
        sin=dp[dp["Estado"].astype(str).str.lower()=="sin registrar"]["Valor pagado"].sum()
        ctx.append(f"PAGOS — Total:{fmt(tp)} | Reg:{fmt(reg)} | Sin:{fmt(sin)}")
    det=df.groupby("Razon Social").agg(
        Saldo=("Saldo Actual","sum"),Max=("Saldo Actual","max"),
        Est=("Estado de cartera",lambda x:x.mode()[0] if not x.mode().empty else ""),
        Int=("Intervalo-Actual",lambda x:x.mode()[0] if not x.mode().empty else ""),
    ).reset_index()
    ctx.append("CLIENTES: "+" || ".join([f"{r['Razon Social']}|{fmt(r['Saldo'])}|{fmt(r['Max'])}|{r['Est']}|{r['Int']}" for _,r in det.iterrows()]))
    return "\n".join(ctx)

def _consultar_ia(pregunta, df, dp):
    ctx=_construir_contexto(df,dp)
    payload={"model":"claude-sonnet-4-20250514","max_tokens":1000,
             "system":"Eres experto en cartera de IPS en Colombia. Responde en español claro y profesional. "
                      "Usa lenguaje del sector salud (EPS,IPS,glosas,ADRES). Máx 200 palabras. Sin HTML.",
             "messages":[{"role":"user","content":f"DATOS:\n{ctx}\n\nPREGUNTA: {pregunta}"}]}
    try:
        r=requests.post("https://api.anthropic.com/v1/messages",
                        headers={"Content-Type":"application/json"},json=payload,timeout=30)
        if r.status_code==200: return r.json()["content"][0]["text"]
        return _respuesta_local(pregunta,df,dp)
    except: return _respuesta_local(pregunta,df,dp)

def _respuesta_local(pregunta, df, dp):
    p=pregunta.lower(); total=df["Saldo Actual"].sum()
    cm=None
    for c in df["Razon Social"].unique():
        if any(w in p for w in [w for w in c.lower().split() if len(w)>3]):
            cm=c; break
    if cm:
        dfc=df[df["Razon Social"]==cm]
        r=(f"Cliente: {cm}\nSaldo: {fmt(dfc['Saldo Actual'].sum())} ({fmt_pct(dfc['Saldo Actual'].sum()/total*100)})\n"
           f"Factura mayor: {fmt(dfc['Saldo Actual'].max())}\n"
           f"Estado: {dfc['Estado de cartera'].mode()[0] if not dfc['Estado de cartera'].mode().empty else 'N/A'}\n"
           f"Intervalo: {dfc['Intervalo-Actual'].mode()[0] if not dfc['Intervalo-Actual'].mode().empty else 'N/A'}\n"
           f"Facturas: {len(dfc):,}")
        return r
    if any(x in p for x in ["mayor","maximo","top","mas alto"]):
        t=df.groupby("Razon Social")["Saldo Actual"].sum()
        return f"Mayor deudor: {t.idxmax()} con {fmt(t.max())} ({fmt_pct(t.max()/total*100)})."
    if "vencida" in p or "mora" in p:
        tv=df["_num_vencida"].sum() if "_num_vencida" in df.columns else 0
        return f"Cartera vencida: {fmt(tv)} ({fmt_pct(tv/total*100)})."
    if "cobro" in p:
        tc=df[df["Estado de cartera"].astype(str).str.strip()=="Para cobro"]["Saldo Actual"].sum()
        return f"Para cobro: {fmt(tc)} ({fmt_pct(tc/total*100)})."
    if "pago" in p and dp is not None:
        tp=dp["Valor pagado"].sum(); sin=dp[dp["Estado"].str.lower()=="sin registrar"]["Valor pagado"].sum()
        return f"Total pagos: {fmt(tp)}. Sin registrar: {fmt(sin)}."
    if "radicar" in p:
        pr=df[df["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()
        return f"Por radicar: {fmt(pr)} ({fmt_pct(pr/total*100)})."
    if "glosa" in p and tiene(df,"glosa_saldo"):
        return f"Saldo glosas: {fmt(df[COLS_OPT['glosa_saldo']].sum())} ({fmt_pct(df[COLS_OPT['glosa_saldo']].sum()/total*100)})."
    return f"Cartera total: {fmt(total)} en {df['Razon Social'].nunique()} clientes. Menciona un cliente o concepto específico."

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — ENVEJECIMIENTO + VENCIMIENTOS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_aging_vencimientos(df):
    section("📊 Envejecimiento y Vencimientos de Cartera")
    total=df["Saldo Actual"].sum()
    ORDEN=["Por radicar","1-30","31-60","61-90","91-150","151-210","211-360",
           "2 años","3 años","4 años","5 años","> 5 AÑOS"]
    COLS_AG=["#2D9B6F","#57C98A","#F4D35E","#F4A261","#E76F51","#E63946",
             "#9B2335","#6B0F1A","#4A0A12","#2D0508","#1A0005","#000000"]
    ag=df.groupby("Intervalo-Actual")["Saldo Actual"].sum().reset_index()
    ag["_o"]=ag["Intervalo-Actual"].apply(
        lambda v:next((i for i,o in enumerate(ORDEN) if str(v).strip().upper()==o.upper()),len(ORDEN)))
    ag=ag.sort_values("_o").drop(columns=["_o"])
    ag["Fmt"]=ag["Saldo Actual"].apply(fmt); ag["% Part"]=(ag["Saldo Actual"]/total*100).round(1)
    ag["Color"]=[COLS_AG[min(i,len(COLS_AG)-1)] for i in range(len(ag))]
    cat=ag["Intervalo-Actual"].tolist()

    fig_aging=go.Figure()
    for _,row in ag.iterrows():
        fig_aging.add_trace(go.Bar(
            x=[row["Saldo Actual"]],y=[row["Intervalo-Actual"]],orientation="h",
            marker_color=row["Color"],text=f"  {row['Fmt']}  ({row['% Part']}%)",
            textposition="outside",name=row["Intervalo-Actual"],showlegend=False))
    fig_aging.update_layout(
        height=420,barmode="stack",paper_bgcolor="white",plot_bgcolor="white",
        font=dict(family="Arial",size=11,color="#333"),margin=dict(l=10,r=150,t=38,b=10),
        xaxis=dict(showgrid=True,gridcolor="#F0F0F0",linecolor="#DEE2E6"),
        yaxis=dict(showgrid=False,linecolor="#DEE2E6",categoryorder="array",
                   categoryarray=list(reversed(cat))),
        title=dict(text="Distribución por Intervalo de Edad",font=dict(size=13,color="#1B3A6B"),x=0))
    plot(fig_aging,use_container_width=True)
    st.dataframe(ag[["Intervalo-Actual","Fmt","% Part"]].rename(
        columns={"Intervalo-Actual":"Intervalo","Fmt":"Saldo"}),use_container_width=True,hide_index=True)

    section("📅 Saldos de Factura Vencida y Por Vencer")
    tv=df["_num_vencida"].sum() if "_num_vencida" in df.columns else 0
    tpv=df["_num_por_vencer"].sum() if "_num_por_vencer" in df.columns else 0
    if tv>0 or tpv>0:
        cv1,cv2,cv3=st.columns(3)
        with cv1: kpi_card("Factura Vencida",fmt(tv),fmt_pct(tv/total*100 if total else 0),"danger")
        with cv2: kpi_card("Por Vencer",fmt(tpv),fmt_pct(tpv/total*100 if total else 0),"warning")
        with cv3: kpi_card("Exposición Total",fmt(tv+tpv),"Vencida + Por vencer","danger")
        col_v="_num_vencida"; col_pv="_num_por_vencer"
        cols_disp=[c for c in [col_v,col_pv] if c in df.columns]
        if cols_disp:
            grp=df.groupby("Razon Social")[cols_disp].sum()
            grp=grp[(grp>0).any(axis=1)]
            grp["_total"]=grp.sum(axis=1)
            grp=grp.nlargest(20,"_total").drop(columns=["_total"]).reset_index()
            fig_stack=go.Figure()
            if "_num_vencida" in grp.columns:
                fig_stack.add_trace(go.Bar(name="Factura Vencida",x=grp["Razon Social"],y=grp["_num_vencida"],
                    marker_color="#E63946",text=grp["_num_vencida"].apply(fmt),textposition="inside",textfont=dict(size=9)))
            if "_num_por_vencer" in grp.columns:
                fig_stack.add_trace(go.Bar(name="Por Vencer",x=grp["Razon Social"],y=grp["_num_por_vencer"],
                    marker_color="#F4A261",text=grp["_num_por_vencer"].apply(fmt),textposition="inside",textfont=dict(size=9)))
            fig_stack.update_layout(barmode="stack",height=420,paper_bgcolor="white",plot_bgcolor="white",
                font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=10,t=40,b=120),
                xaxis=dict(tickangle=-35,showgrid=False,linecolor="#DEE2E6"),
                yaxis=dict(gridcolor="#F0F0F0",linecolor="#DEE2E6"),
                legend=dict(orientation="h",yanchor="bottom",y=-0.45,xanchor="center",x=0.5),
                title=dict(text="Top 20 Clientes — Vencida vs Por Vencer",font=dict(size=13,color="#1B3A6B"),x=0))
            plot(fig_stack,use_container_width=True)

    section("🚨 Alertas Proactivas por Intervalo")
    RIESGO={"1-30":("🔴","CRÍTICO","próximos 30 días"),"31-60":("🟠","ALTO","31-60 días"),
            "61-90":("🟡","MEDIO","61-90 días"),"91-150":("🟡","MEDIO","91-150 días")}
    ok=False
    for interv,(icono,nivel,desc) in RIESGO.items():
        di=df[df["Intervalo-Actual"].astype(str).str.strip()==interv]
        if di.empty: continue
        si=di["Saldo Actual"].sum(); pi=(si/total*100) if total else 0
        top=di.groupby("Razon Social")["Saldo Actual"].sum().nlargest(3)
        cli=", ".join([f"{c} ({fmt(v)})" for c,v in top.items()])
        kind="danger" if nivel=="CRÍTICO" else "warning"
        fref=""
        if tiene(df,"fecha_venc"):
            mu=di[COLS_OPT["fecha_venc"]].dropna()
            if not mu.empty: fref=" Validar en columna 'Fecha de Vencimiento'."
        msg=(f"{icono} <b>Riesgo {nivel} — {interv} ({desc}):</b> {fmt(si)} ({fmt_pct(pi)}). "
             f"Clientes: {cli}.{fref}")
        st.markdown(f'<div class="insight-box {kind}">{msg}</div>',unsafe_allow_html=True)
        ok=True
    if not ok: st.info("No se encontraron registros en intervalos de riesgo próximo.")

    section("👥 Relación de Clientes por Intervalo")
    isel=st.selectbox("Selecciona un intervalo:",["(Todos)"]+cat)
    df_f=df if isel=="(Todos)" else df[df["Intervalo-Actual"]==isel]
    ca=(df_f.groupby(["Razon Social","Intervalo-Actual","Tipo De Empresa","Estado de cartera"])
        ["Saldo Actual"].sum().reset_index().sort_values("Saldo Actual",ascending=False))
    ca["% Part"]=(ca["Saldo Actual"]/total*100).round(2)
    ca["Saldo Actual"]=ca["Saldo Actual"].apply(fmt)
    st.dataframe(ca.rename(columns={"Razon Social":"Cliente","Saldo Actual":"Saldo",
                                     "Tipo De Empresa":"Tipo","Estado de cartera":"Estado"}),
                 use_container_width=True,hide_index=True,height=400)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — TIPO DE EMPRESA
# ══════════════════════════════════════════════════════════════════════════════
def seccion_tipo_empresa(df):
    section("🏢 Análisis por Tipo de Empresa")
    total=df["Saldo Actual"].sum()
    t=(df.groupby("Tipo De Empresa")["Saldo Actual"].sum()
       .reset_index().sort_values("Saldo Actual",ascending=False))
    t["Saldo"]=t["Saldo Actual"].apply(fmt); t["% Part."]=(t["Saldo Actual"]/total*100).round(1)
    c1,c2=st.columns([1.3,1])
    with c1:
        fig_te=go.Figure()
        max_v=t["Saldo Actual"].max()
        for i,row in t.iterrows():
            intens=row["Saldo Actual"]/max_v if max_v else 0
            rr=int(27+(180-27)*intens); gg=int(58+(90-58)*(1-intens)); bb=107
            fig_te.add_trace(go.Bar(x=[row["Saldo Actual"]],y=[row["Tipo De Empresa"]],
                orientation="h",marker=dict(color=f"rgb({rr},{gg},{bb})",line=dict(color="white",width=1)),
                text=f"  {row['Saldo']}  ({row['% Part.']}%)",textposition="outside",
                name=row["Tipo De Empresa"],showlegend=False))
        fig_te.update_layout(height=380,barmode="overlay",paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=11,color="#333"),margin=dict(l=10,r=160,t=38,b=10),
            xaxis=dict(showgrid=True,gridcolor="#F0F0F0",linecolor="#DEE2E6"),
            yaxis=dict(showgrid=False,linecolor="#DEE2E6",categoryorder="array",
                       categoryarray=list(reversed(t["Tipo De Empresa"].tolist()))),
            title=dict(text="Cartera por Tipo de Empresa",font=dict(size=13,color="#1B3A6B"),x=0))
        plot(fig_te,use_container_width=True)
    with c2:
        fig_do=go.Figure(go.Pie(labels=t["Tipo De Empresa"],values=t["Saldo Actual"],hole=0.5,
            textinfo="percent",textposition="inside",
            hovertemplate="<b>%{label}</b><br>Saldo: %{customdata}<br>Part: %{percent}<extra></extra>",
            customdata=t["Saldo"],marker=dict(colors=PALETTE[:len(t)],line=dict(color="white",width=2))))
        fig_do.update_layout(height=380,paper_bgcolor="white",font=dict(family="Arial",size=11,color="#333"),
            legend=dict(orientation="v",x=1.02,y=0.5,xanchor="left"),margin=dict(l=10,r=10,t=38,b=10),
            title=dict(text="Participación %",font=dict(size=13,color="#1B3A6B"),x=0),
            annotations=[dict(text=f"<b>{fmt(total)}</b>",x=0.5,y=0.5,font=dict(size=11,color="#1B3A6B"),showarrow=False)])
        plot(fig_do,use_container_width=True)
    st.dataframe(t[["Tipo De Empresa","Saldo","% Part."]].rename(columns={"Tipo De Empresa":"Tipo de Empresa"}),
                 use_container_width=True,hide_index=True)
    section("📋 Análisis Interpretativo por Tipo de Empresa")
    st.caption(f"Análisis sobre {len(t)} tipos con {fmt(total)} en cartera.")
    INTERP={
        "eps":("EPS — Entidades Promotoras de Salud","Principal pagador del sistema. Alta rotación pero con frecuentes glosas. Gestionar con tesorería y auditoría médica de la EPS. Revisar acuerdos de pago y procesos jurídicos activos.","warning"),
        "alcaldia":("Alcaldías","Tiempos prolongados por procesos presupuestales y PAC. Riesgo ante cambios de administración. Radicar antes del cierre fiscal.","warning"),
        "gobernacion":("Gobernaciones","Sujeto a ejecución presupuestal departamental. Alta dependencia de giros ADRES. Gestionar con secretarías de salud.","warning"),
        "particular":("Particulares","Mayor riesgo de irrecuperabilidad. Activar cobro persuasivo y jurídico temprano. Proponer acuerdos de pago en cuotas.","danger"),
        "liquidaci":("En Liquidación","Recuperación incierta. Activar acreencias ante el liquidador. Requiere provisión contable inmediata.","danger"),
        "soat":("SOAT / Aseguradoras","Respaldado por póliza. Verificar vigencia y radicación completa de soportes.","info"),
    }
    for _,row in t.iterrows():
        tl=str(row["Tipo De Empresa"]).lower()
        ik=next((k for k in INTERP if k in tl),None)
        titulo,txt,kind=INTERP.get(ik,(row["Tipo De Empresa"],
            f"Representa el {row['% Part.']}% ({row['Saldo']}). Revisar mora y gestionar acuerdos.","info"))
        df_tipo=df[df["Tipo De Empresa"]==row["Tipo De Empresa"]]
        n_cli_t=df_tipo["Razon Social"].nunique()
        ep=df_tipo.groupby("Estado de cartera")["Saldo Actual"].sum().idxmax() if not df_tipo.empty else "N/A"
        pct_c_t=(df_tipo[df_tipo["Estado de cartera"].astype(str).str.strip()=="Para cobro"]["Saldo Actual"].sum()
                /row["Saldo Actual"]*100) if row["Saldo Actual"]>0 else 0
        icono="⚠️" if kind=="danger" else "📌"
        st.markdown(f"""<div class="analisis-box">
        <b>{icono} {row['Tipo De Empresa']}</b> — Saldo: <b>{row['Saldo']}</b> | Part: <b>{row['% Part.']}%</b>
        | Clientes: <b>{n_cli_t}</b> | Estado ppal: <b>{ep}</b> | % Para cobro: <b>{pct_c_t:.1f}%</b><br><br>
        {txt}</div>""",unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — GLOSAS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_glosas(df):
    section("🔍 Análisis de Glosas")
    if not tiene(df,"glosa_valor") and not tiene(df,"glosa_saldo"):
        st.info("No se encontraron columnas de glosas."); return
    cg=COLS_OPT["glosa_valor"] if tiene(df,"glosa_valor") else None
    cgs=COLS_OPT["glosa_saldo"] if tiene(df,"glosa_saldo") else None
    dfg=df[(df[cg]>0) if cg else pd.Series([True]*len(df),index=df.index)].copy()
    if dfg.empty: st.info("No hay registros con glosas."); return
    total=df["Saldo Actual"].sum()
    tg=dfg[cgs].sum() if cgs else (dfg[cg].sum() if cg else 0)
    tv=dfg[cg].sum() if cg else 0; pct=(tg/total*100) if total else 0
    col_est=COLS_OPT.get("glosa_estado"); estados_ok=["en acta de conciliacion","conciliada","en acta","conciliado"]
    alerta_2026=pd.DataFrame()
    col_rec=COLS_OPT.get("recepcion"); col_fac=COLS_OPT.get("fecha_fac"); fcu=None
    for cf in [col_rec,col_fac]:
        if cf and cf in dfg.columns:
            dfg["_ft"]=pd.to_datetime(dfg[cf],errors="coerce")
            if (dfg["_ft"].dt.year==2026).any(): fcu=cf; break
    if fcu and col_est and col_est in dfg.columns:
        dfg["_ft"]=pd.to_datetime(dfg[fcu],errors="coerce")
        d26=dfg[dfg["_ft"].dt.year==2026].copy()
        alerta_2026=d26[~d26[col_est].astype(str).str.strip().str.lower().isin(estados_ok)]
    if not alerta_2026.empty:
        va=alerta_2026[cgs].sum() if cgs else alerta_2026[cg].sum() if cg else 0
        st.error(f"🚨 Glosas 2026 sin conciliar: {len(alerta_2026):,} registros | {fmt(va)} | {alerta_2026['Razon Social'].nunique()} clientes")
        with st.expander("👁️ Ver detalle glosas 2026"):
            ca=["Razon Social","Tipo De Empresa",col_est,fcu,"Intervalo-Actual"]
            if cg and cg in alerta_2026.columns: ca.append(cg)
            if cgs and cgs in alerta_2026.columns: ca.append(cgs)
            da=alerta_2026[[c for c in ca if c in alerta_2026.columns]].copy()
            for nc in ([cg] if cg else [])+([cgs] if cgs else []):
                if nc in da.columns: da[nc]=da[nc].apply(fmt)
            st.dataframe(da,use_container_width=True,hide_index=True)
            cv=cgs if cgs else cg
            if cv:
                alerta_2026[cv]=pd.to_numeric(alerta_2026[cv],errors="coerce").fillna(0)
                tea=alerta_2026.groupby("Tipo De Empresa").agg(Registros=(cv,"count"),Valor=(cv,"sum")).reset_index().sort_values("Valor",ascending=False)
                tea["Valor"]=tea["Valor"].apply(fmt)
                st.markdown("**Por Tipo de Empresa:**"); st.dataframe(tea,use_container_width=True,hide_index=True)
    c1,c2,c3,c4=st.columns(4)
    with c1: kpi_card("Saldo Glosas s/Cartera",fmt(tg),f"{len(dfg):,} registros","warning")
    with c2: kpi_card("Vr Glosa Reportada",fmt(tv),"Valor bruto","warning")
    with c3: kpi_card("% sobre Cartera",fmt_pct(pct),"Impacto financiero","warning")
    with c4:
        if tiene(df,"glosa_dias"): kpi_card("Días Glosa Prom.",f"{dfg[COLS_OPT['glosa_dias']].mean():.0f} días","","warning")
        else: kpi_card("Clientes c/Glosa",str(dfg["Razon Social"].nunique()),"afectados","warning")
    cl1,cl2=st.columns(2)
    with cl1:
        if col_est and col_est in dfg.columns:
            vc=cgs if cgs else cg
            if vc:
                gx=dfg.groupby(col_est)[vc].sum().reset_index(); gx.columns=["Estado","Valor"]; gx["Fmt"]=gx["Valor"].apply(fmt)
                gx=gx.sort_values("Valor",ascending=False)
                ce=["#2D9B6F" if any(x in str(e).lower() for x in ["conciliada","acta","cerrada"])
                    else ("#F4A261" if "respuesta" in str(e).lower() else "#E63946") for e in gx["Estado"]]
                fg=go.Figure(go.Bar(x=gx["Valor"],y=gx["Estado"],orientation="h",
                    marker_color=ce,text=gx["Fmt"],textposition="outside"))
                fg.update_layout(height=340,paper_bgcolor="white",plot_bgcolor="white",
                    font=dict(family="Arial",size=10),margin=dict(l=10,r=140,t=38,b=10),
                    xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),yaxis=dict(showgrid=False),
                    title=dict(text="Glosas por Estado",font=dict(size=12,color="#1B3A6B"),x=0),showlegend=False)
                plot(fg,use_container_width=True)
    with cl2:
        if tiene(df,"glosa_intervalo"):
            vc=cgs if cgs else cg
            if vc:
                gi=dfg.groupby(COLS_OPT["glosa_intervalo"])[vc].sum().reset_index(); gi.columns=["Intervalo","Valor"]; gi["Fmt"]=gi["Valor"].apply(fmt)
                gi=gi.sort_values("Valor",ascending=False)
                fg2=go.Figure(go.Bar(x=gi["Intervalo"],y=gi["Valor"],
                    marker=dict(color=gi["Valor"],colorscale=[[0,"#2D9B6F"],[0.5,"#F4A261"],[1,"#E63946"]],showscale=True,
                                colorbar=dict(title="Valor",thickness=12,len=0.7)),text=gi["Fmt"],textposition="outside"))
                fg2.update_layout(height=340,paper_bgcolor="white",plot_bgcolor="white",
                    font=dict(family="Arial",size=10),margin=dict(l=10,r=10,t=38,b=10),
                    xaxis=dict(showgrid=False),yaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
                    title=dict(text="Glosas por Intervalo",font=dict(size=12,color="#1B3A6B"),x=0),showlegend=False)
                plot(fg2,use_container_width=True)
    section("📋 Resumen por Cliente")
    ad={}
    if cg and cg in dfg.columns: ad[cg]="sum"
    if cgs and cgs in dfg.columns: ad[cgs]="sum"
    if ad:
        tg2=dfg.groupby("Razon Social").agg(ad).reset_index().sort_values(list(ad.keys())[0],ascending=False)
        for c_ in ad: tg2[c_]=tg2[c_].apply(fmt)
        st.dataframe(tg2,use_container_width=True,hide_index=True)
    section("📄 Detalle de Registros")
    cd=["Razon Social","Tipo De Empresa","Intervalo-Actual","Estado de cartera",
        "Vr,Glosa Reportada","R-AUDITORIA","Dias de Glosa","Intervalo Dias Glosa",
        "Estado De Recepcion","Estado De Glosa","Fecha_Recepcion",
        "Saldo de Glosa según cartera","Devuelta","Saldo Actual","OBSERVACION"]
    det=dfg[[c for c in cd if c in dfg.columns]].copy()
    for nc in ["Vr,Glosa Reportada","Saldo de Glosa según cartera","Saldo Actual"]:
        if nc in det.columns:
            det[nc]=pd.to_numeric(det[nc],errors="coerce").apply(lambda x:fmt(x) if pd.notna(x) else "")
    st.dataframe(det,use_container_width=True,hide_index=True,height=400)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — ÍNDICE DE GLOSAS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_indice_glosas(df):
    section("📊 Índice de Glosas por Cliente")
    col_vf=COLS_OPT.get("vr_factura"); col_vg=COLS_OPT.get("glosa_valor"); col_vgs=COLS_OPT.get("glosa_saldo")
    if not (col_vf and col_vf in df.columns): st.info("No se encontró 'Vr factura'."); return
    dfi=df[df[col_vf]>0].copy()
    if dfi.empty: st.info("No hay registros con valor de factura."); return
    vg_col=col_vgs if (col_vgs and col_vgs in dfi.columns) else (col_vg if (col_vg and col_vg in dfi.columns) else None)
    if not vg_col: st.info("No se encontró columna de glosa."); return
    total_fac=dfi[col_vf].sum(); total_glo=dfi[vg_col].sum()
    idx_global=(total_glo/total_fac*100) if total_fac else 0
    c1,c2,c3,c4=st.columns(4)
    with c1: kpi_card("Total Facturado",fmt(total_fac),"Base de cálculo","primary")
    with c2: kpi_card("Total Glosado",fmt(total_glo),"Valor bruto glosado","warning")
    with c3: kpi_card("Índice de Glosas",f"{idx_global:.2f}%","Glosado/Facturado","danger" if idx_global>5 else "warning")
    with c4:
        if col_vgs and col_vgs in dfi.columns:
            neto=dfi[col_vf].sum()-dfi[col_vgs].sum()
            kpi_card("Neto Cobrable",fmt(neto),"Facturado menos glosas","success")
    st.markdown("""<div class="insight-box warning">📌 <b>Referencia sector salud Colombia:</b>
    Índice saludable: <b>menor al 5%</b>. Entre 5%-10%: alerta. Mayor al 10%: intervención inmediata en auditoría médica.</div>""",
    unsafe_allow_html=True)
    section("📋 Índice por Cliente (Top 15)")
    grp=dfi.groupby("Razon Social").agg(Facturado=(col_vf,"sum"),Glosado=(vg_col,"sum")).reset_index()
    grp=grp[grp["Facturado"]>0].copy()
    grp["Índice %"]=(grp["Glosado"]/grp["Facturado"]*100).round(2)
    grp=grp.sort_values("Índice %",ascending=False).head(15)
    grp["Facturado_fmt"]=grp["Facturado"].apply(fmt); grp["Glosado_fmt"]=grp["Glosado"].apply(fmt)
    ci=[("#E63946" if v>10 else ("#F4A261" if v>5 else "#2D9B6F")) for v in grp["Índice %"]]
    fig_ig=go.Figure(go.Bar(x=grp["Índice %"],y=grp["Razon Social"],orientation="h",marker_color=ci,
        text=[f"{v:.1f}%" for v in grp["Índice %"]],textposition="outside",
        customdata=list(zip(grp["Facturado_fmt"],grp["Glosado_fmt"])),
        hovertemplate="<b>%{y}</b><br>Facturado: %{customdata[0]}<br>Glosado: %{customdata[1]}<br>Índice: %{x:.2f}%<extra></extra>"))
    fig_ig.add_vline(x=5,line_dash="dash",line_color="#F4A261",annotation_text="Alerta 5%")
    fig_ig.add_vline(x=10,line_dash="dash",line_color="#E63946",annotation_text="Crítico 10%")
    fig_ig.update_layout(height=460,paper_bgcolor="white",plot_bgcolor="white",
        font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=80,t=38,b=10),
        xaxis=dict(showgrid=True,gridcolor="#F0F0F0",title="Índice de Glosas (%)"),
        yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(grp["Razon Social"].tolist()))),
        title=dict(text="Índice de Glosas (🔴>10% 🟡5-10% 🟢<5%)",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
    plot(fig_ig,use_container_width=True)
    st.dataframe(grp[["Razon Social","Facturado_fmt","Glosado_fmt","Índice %"]].rename(
        columns={"Razon Social":"Cliente","Facturado_fmt":"Facturado","Glosado_fmt":"Glosado"}),
        use_container_width=True,hide_index=True)
    section("🏢 Índice por Tipo de Empresa")
    grp_t=dfi.groupby("Tipo De Empresa").agg(Facturado=(col_vf,"sum"),Glosado=(vg_col,"sum")).reset_index()
    grp_t=grp_t[grp_t["Facturado"]>0].copy()
    grp_t["Índice %"]=(grp_t["Glosado"]/grp_t["Facturado"]*100).round(2)
    grp_t=grp_t.sort_values("Índice %",ascending=False)
    grp_t["Facturado"]=grp_t["Facturado"].apply(fmt); grp_t["Glosado"]=grp_t["Glosado"].apply(fmt)
    st.dataframe(grp_t.rename(columns={"Tipo De Empresa":"Tipo"}),use_container_width=True,hide_index=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 7 — RADICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
def seccion_radicacion(df):
    section("⏱️ Eficiencia Operativa de Radicación")
    if "_dias_rad" not in df.columns: st.info("No se encontraron columnas de fechas de radicación."); return
    dfi=df[df["_dias_rad"].notna()&(df["_dias_rad"]>=0)].copy()
    if dfi.empty: st.info("Sin datos de radicación."); return
    prom_dias=dfi["_dias_rad"].mean(); mediana=dfi["_dias_rad"].median(); max_dias=dfi["_dias_rad"].max()
    dentro_obj=0; pct_dentro=0
    if "_dentro_objetivo" in dfi.columns:
        dentro_obj=dfi["_dentro_objetivo"].sum(); pct_dentro=(dentro_obj/len(dfi)*100)
    c1,c2,c3,c4=st.columns(4)
    with c1: kpi_card("Días Prom. Radicación",f"{prom_dias:.0f} días","Desde fecha factura","warning" if prom_dias>30 else "success")
    with c2: kpi_card("Mediana",f"{mediana:.0f} días","50% de facturas","default")
    with c3: kpi_card("Máximo",f"{max_dias:.0f} días","caso más demorado","danger" if max_dias>90 else "warning")
    with c4:
        if pct_dentro>0: kpi_card("Dentro del Objetivo",fmt_pct(pct_dentro),f"{int(dentro_obj):,} facturas","success" if pct_dentro>80 else "warning")
        else: kpi_card("Total Analizadas",f"{len(dfi):,}","facturas","default")
    st.markdown("""<div class="insight-box info">📌 <b>Referencia:</b> Promedio mayor a <b>30 días</b> indica ineficiencia operativa
    que puede retrasar el inicio del plazo de pago del pagador.</div>""",unsafe_allow_html=True)
    col1,col2=st.columns(2)
    with col1:
        dfi_hist=dfi[dfi["_dias_rad"]<=365].copy()
        fig_hist=px.histogram(dfi_hist,x="_dias_rad",nbins=30,color_discrete_sequence=["#2E86AB"],
            labels={"_dias_rad":"Días para radicar","count":"Facturas"})
        fig_hist.add_vline(x=30,line_dash="dash",line_color="#F4A261",annotation_text="30 días")
        fig_hist.add_vline(x=prom_dias,line_dash="dot",line_color="#E63946",annotation_text=f"Prom:{prom_dias:.0f}d")
        plot(fig_layout(fig_hist,"Distribución de Días de Radicación",360),use_container_width=True)
    with col2:
        bins=[0,15,30,60,90,180,365,9999]; lbls=["0-15","16-30","31-60","61-90","91-180","181-365","+365"]
        dfi["_rr"]=pd.cut(dfi["_dias_rad"],bins=bins,labels=lbls,right=True)
        rng=dfi.groupby("_rr",observed=True)["Saldo Actual"].agg(["sum","count"]).reset_index()
        rng.columns=["Rango","Saldo","Facturas"]; rng["Fmt"]=rng["Saldo"].apply(fmt)
        COLS_R=["#2D9B6F","#57C98A","#F4D35E","#F4A261","#E76F51","#E63946","#7B1D1D"]
        fig_rng=go.Figure(go.Bar(x=rng["Rango"],y=rng["Saldo"],marker_color=COLS_R[:len(rng)],
            text=rng["Fmt"],textposition="outside",customdata=rng["Facturas"],
            hovertemplate="Rango: %{x}<br>Saldo: %{text}<br>Facturas: %{customdata}<extra></extra>"))
        plot(fig_layout(fig_rng,"Saldo por Rango de Días",360),use_container_width=True)
    section("🐢 Top 15 Clientes con Mayor Demora")
    top_dem=dfi.groupby("Razon Social")["_dias_rad"].mean().reset_index()
    top_dem.columns=["Cliente","Días Promedio"]; top_dem=top_dem.sort_values("Días Promedio",ascending=False).head(15)
    top_dem["Días Promedio"]=top_dem["Días Promedio"].round(0).astype(int)
    fig_dem=go.Figure(go.Bar(x=top_dem["Días Promedio"],y=top_dem["Cliente"],orientation="h",
        marker_color=["#E63946" if v>60 else ("#F4A261" if v>30 else "#2D9B6F") for v in top_dem["Días Promedio"]],
        text=top_dem["Días Promedio"].astype(str)+" días",textposition="outside"))
    fig_dem.update_layout(height=440,paper_bgcolor="white",plot_bgcolor="white",
        font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=100,t=38,b=10),
        xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
        yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(top_dem["Cliente"].tolist()))),
        title=dict(text="Días Promedio de Radicación por Cliente",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
    plot(fig_dem,use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 8 — DEVUELTAS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_devueltas(df):
    section("↩️ Análisis de Facturas Devueltas")
    col_dev=COLS_OPT.get("devuelta"); col_devj=COLS_OPT.get("dev_juridica"); col_mot=COLS_OPT.get("motivo_dev")
    mask_dev=pd.Series([False]*len(df),index=df.index)
    if col_dev  and col_dev  in df.columns: mask_dev=mask_dev|df[col_dev].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])
    if col_devj and col_devj in df.columns: mask_dev=mask_dev|df[col_devj].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])
    dfd=df[mask_dev].copy(); total=df["Saldo Actual"].sum()
    if dfd.empty: st.info("No se identificaron facturas devueltas."); return
    t_dev=dfd["Saldo Actual"].sum(); pct_dev=(t_dev/total*100) if total else 0
    c1,c2,c3,c4=st.columns(4)
    with c1: kpi_card("Total Devuelto",fmt(t_dev),fmt_pct(pct_dev),"danger")
    with c2: kpi_card("Facturas Devueltas",f"{len(dfd):,}","registros","warning")
    with c3: kpi_card("Clientes Afectados",f"{dfd['Razon Social'].nunique():,}","clientes","warning")
    with c4: kpi_card("Promedio x Factura",fmt(t_dev/len(dfd) if len(dfd)>0 else 0),"valor promedio","default")
    col1,col2=st.columns(2)
    with col1:
        top_dev=dfd.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10).reset_index()
        top_dev.columns=["Cliente","Saldo"]; top_dev["Fmt"]=top_dev["Saldo"].apply(fmt)
        fig_dc=go.Figure(go.Bar(x=top_dev["Saldo"],y=top_dev["Cliente"],orientation="h",
            marker_color="#E63946",text=top_dev["Fmt"],textposition="outside"))
        fig_dc.update_layout(height=380,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=140,t=38,b=10),
            xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
            yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(top_dev["Cliente"].tolist()))),
            title=dict(text="Top 10 Clientes con Más Devoluciones",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
        plot(fig_dc,use_container_width=True)
    with col2:
        if col_mot and col_mot in dfd.columns:
            mot=dfd.groupby(col_mot)["Saldo Actual"].sum().reset_index()
            mot.columns=["Motivo","Saldo"]
            mot=mot[~mot["Motivo"].astype(str).str.strip().str.lower().isin(["","nan"])].sort_values("Saldo",ascending=False).head(10)
            mot["Fmt"]=mot["Saldo"].apply(fmt)
            fig_mot=go.Figure(go.Bar(x=mot["Saldo"],y=mot["Motivo"],orientation="h",
                marker_color=PALETTE[:len(mot)],text=mot["Fmt"],textposition="outside"))
            fig_mot.update_layout(height=380,paper_bgcolor="white",plot_bgcolor="white",
                font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=140,t=38,b=10),
                xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
                yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(mot["Motivo"].tolist()))),
                title=dict(text="Por Motivo de Devolución",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
            plot(fig_mot,use_container_width=True)
    section("📄 Detalle")
    cols_d=["Razon Social","Tipo De Empresa","Intervalo-Actual","Estado de cartera","Saldo Actual"]
    if col_mot  and col_mot  in dfd.columns: cols_d.append(col_mot)
    if "OBSERVACION" in dfd.columns: cols_d.append("OBSERVACION")
    det=dfd[[c for c in cols_d if c in dfd.columns]].copy()
    det["Saldo Actual"]=det["Saldo Actual"].apply(fmt)
    st.dataframe(det.sort_values("Saldo Actual",ascending=False),use_container_width=True,hide_index=True,height=360)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 9 — ACUERDOS Y PGP
# ══════════════════════════════════════════════════════════════════════════════
def seccion_acuerdos_pgp(df):
    section("🤝 Panel de Acuerdos de Pago y PGP")
    total=df["Saldo Actual"].sum()
    col_ss=COLS_OPT.get("supersalud"); col_vra=COLS_OPT.get("vr_acuerdo"); col_pgp=COLS_OPT.get("pgp")
    if col_vra and col_vra in df.columns:
        df_ac=df[df[col_vra]>0].copy()
        if not df_ac.empty:
            t_ac=df_ac["Saldo Actual"].sum(); vr_ac=df_ac[col_vra].sum()
            c1,c2,c3,c4=st.columns(4)
            with c1: kpi_card("Cartera con Acuerdo",fmt(t_ac),fmt_pct(t_ac/total*100 if total else 0),"success")
            with c2: kpi_card("Valor Acordado",fmt(vr_ac),"Monto comprometido","success")
            with c3: kpi_card("Clientes con Acuerdo",str(df_ac["Razon Social"].nunique()),"clientes","primary")
            with c4: kpi_card("% Acuerdo vs Saldo",fmt_pct(vr_ac/t_ac*100 if t_ac else 0),"Acordado/saldo","success")
            agg_ac=df_ac.groupby("Razon Social").agg(Saldo_Cartera=("Saldo Actual","sum")).reset_index()
            agg_vra=df_ac.groupby("Razon Social")[col_vra].sum().reset_index(); agg_vra.columns=["Razon Social","Valor Acordado"]
            agg_ac=pd.merge(agg_ac,agg_vra,on="Razon Social",how="left")
            agg_ac["% Acuerdo"]=(agg_ac["Valor Acordado"]/agg_ac["Saldo_Cartera"]*100).round(1)
            agg_ac=agg_ac.sort_values("Saldo_Cartera",ascending=False)
            agg_ac["Saldo_Cartera"]=agg_ac["Saldo_Cartera"].apply(fmt); agg_ac["Valor Acordado"]=agg_ac["Valor Acordado"].apply(fmt)
            st.dataframe(agg_ac.rename(columns={"Razon Social":"Cliente","Saldo_Cartera":"Saldo Cartera"}),
                         use_container_width=True,hide_index=True)
        else: st.info("No se identificaron registros con acuerdos de pago.")
    else: st.info("No se encontró la columna 'VR ACUERDO'.")
    st.markdown("---")
    section("🏥 Cartera PGP vs Por Evento")
    if col_pgp and col_pgp in df.columns:
        df["_es_pgp"]=df[col_pgp].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])
        pgp_df=df[df["_es_pgp"]]; evt_df=df[~df["_es_pgp"]]
        t_pgp=pgp_df["Saldo Actual"].sum(); t_evt=evt_df["Saldo Actual"].sum()
        c1,c2,c3,c4=st.columns(4)
        with c1: kpi_card("Cartera PGP",fmt(t_pgp),fmt_pct(t_pgp/total*100 if total else 0),"primary")
        with c2: kpi_card("Cartera Evento",fmt(t_evt),fmt_pct(t_evt/total*100 if total else 0),"secondary")
        with c3: kpi_card("Clientes PGP",str(pgp_df["Razon Social"].nunique()),"clientes","primary")
        with c4: kpi_card("Clientes Evento",str(evt_df["Razon Social"].nunique()),"clientes","default")
        c1,c2=st.columns(2)
        with c1:
            fig_pgp=go.Figure(go.Pie(labels=["PGP","Por Evento"],values=[t_pgp,t_evt],hole=0.5,
                marker=dict(colors=["#1B3A6B","#2E86AB"],line=dict(color="white",width=2)),
                textinfo="percent+label",textposition="inside"))
            fig_pgp.update_layout(height=320,paper_bgcolor="white",font=dict(family="Arial",size=11,color="#333"),
                margin=dict(l=10,r=10,t=38,b=10),
                title=dict(text="PGP vs Evento",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
            plot(fig_pgp,use_container_width=True)
        with c2:
            pgp_te=pgp_df.groupby("Tipo De Empresa")["Saldo Actual"].sum().nlargest(8).reset_index()
            pgp_te["Fmt"]=pgp_te["Saldo Actual"].apply(fmt)
            fig_pte=px.bar(pgp_te,x="Tipo De Empresa",y="Saldo Actual",color="Tipo De Empresa",
                           color_discrete_sequence=PALETTE,text="Fmt")
            fig_pte.update_traces(textposition="outside")
            plot(fig_layout(fig_pte,"PGP por Tipo de Empresa",320),use_container_width=True)
    else: st.info("No se encontró la columna 'PGP'.")

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 10 — MAPA DE CALOR
# ══════════════════════════════════════════════════════════════════════════════
def seccion_mapa_calor(df):
    section("🗺️ Mapa de Calor — Mora por Cliente e Intervalo")
    ORDEN=["Por radicar","1-30","31-60","61-90","91-150","151-210","211-360",
           "2 años","3 años","4 años","5 años","> 5 AÑOS"]
    top20=df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(20).index.tolist()
    df_mc=df[df["Razon Social"].isin(top20)].copy()
    pivot=df_mc.pivot_table(index="Razon Social",columns="Intervalo-Actual",values="Saldo Actual",aggfunc="sum",fill_value=0)
    cols_ord=[c for c in ORDEN if c in pivot.columns]+[c for c in pivot.columns if c not in ORDEN]
    pivot=pivot[cols_ord]
    pivot.index=[str(i)[:30]+"…" if len(str(i))>30 else str(i) for i in pivot.index]
    fig_hm=go.Figure(go.Heatmap(
        z=pivot.values,x=list(pivot.columns),y=list(pivot.index),
        colorscale=[[0.0,"#EDFDF6"],[0.1,"#A8DADC"],[0.3,"#F4D35E"],[0.6,"#F4A261"],[0.85,"#E63946"],[1.0,"#7B1D1D"]],
        text=[[fmt(v) if v>0 else "" for v in row] for row in pivot.values],
        texttemplate="%{text}",textfont=dict(size=8),
        hovertemplate="<b>%{y}</b><br>Intervalo: %{x}<br>Saldo: %{text}<extra></extra>",xgap=2,ygap=2))
    fig_hm.update_layout(
        height=max(380,len(pivot)*32+80),paper_bgcolor="white",plot_bgcolor="white",
        font=dict(family="Arial",size=9,color="#333"),margin=dict(l=10,r=10,t=50,b=80),
        xaxis=dict(side="top",tickangle=-30,showgrid=False),yaxis=dict(showgrid=False,autorange="reversed"),
        title=dict(text="Saldo por Cliente × Intervalo (Top 20)",font=dict(size=13,color="#1B3A6B"),x=0),
        coloraxis_colorbar=dict(title="Saldo",thickness=12,len=0.6))
    plot(fig_hm,use_container_width=True)
    st.caption("🔴 Rojo oscuro = mayor concentración | 🟢 Verde = saldo bajo o cero")
    with st.expander("📄 Ver tabla cruzada completa"):
        pivot_fmt=pivot.map(lambda x:fmt(x) if x>0 else "-")
        st.dataframe(pivot_fmt,use_container_width=True,height=420)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 11 — COSECHA
# ══════════════════════════════════════════════════════════════════════════════
def seccion_cosecha(df):
    section("📅 Análisis de Cosecha de Facturación por Año")
    total=df["Saldo Actual"].sum()
    col_af=COLS_OPT.get("año_fac"); col_ar=COLS_OPT.get("año_rad")
    if not (col_af and col_af in df.columns): st.info("No se encontró 'Año Facturacion'."); return
    cos=df.groupby(col_af).agg(Saldo=("Saldo Actual","sum"),Facturas=("Saldo Actual","count")).reset_index()
    cos.columns=["Año","Saldo","Facturas"]; cos["Año"]=cos["Año"].astype(str)
    cos["% Part"]=(cos["Saldo"]/total*100).round(1); cos["Saldo Fmt"]=cos["Saldo"].apply(fmt)
    cos=cos.sort_values("Año")
    c1,c2,c3=st.columns(3)
    with c1: kpi_card("Años en Cartera",str(len(cos)),"cosechas distintas","primary")
    with c2: kpi_card("Año Mayor Saldo",cos.loc[cos["Saldo"].idxmax(),"Año"],"mayor acumulación","warning")
    with c3:
        rec=cos[cos["Año"].isin(["2025","2026"])]["Saldo"].sum()
        kpi_card("Cartera 2025-2026",fmt(rec),fmt_pct(rec/total*100 if total else 0),"success")
    col1,col2=st.columns(2)
    with col1:
        fig_cos=go.Figure(go.Bar(x=cos["Año"],y=cos["Saldo"],
            marker=dict(color=cos["Saldo"],colorscale=[[0,"#2D9B6F"],[0.5,"#F4A261"],[1,"#E63946"]],showscale=False),
            text=cos["Saldo Fmt"],textposition="outside",customdata=cos["Facturas"],
            hovertemplate="Año: %{x}<br>Saldo: %{text}<br>Facturas: %{customdata}<extra></extra>"))
        fig_cos.update_layout(height=360,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=11,color="#333"),margin=dict(l=10,r=10,t=38,b=10),
            xaxis=dict(showgrid=False,type="category"),yaxis=dict(gridcolor="#F0F0F0"),
            title=dict(text="Saldo por Año de Facturación",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
        plot(fig_cos,use_container_width=True)
    with col2:
        cos_s=cos.copy(); cos_s["Acumulado"]=cos_s["Saldo"].cumsum()
        fig_acu=go.Figure()
        fig_acu.add_trace(go.Bar(x=cos_s["Año"],y=cos_s["Saldo"],name="Saldo año",marker_color="#2E86AB",opacity=0.6))
        fig_acu.add_trace(go.Scatter(x=cos_s["Año"],y=cos_s["Acumulado"],name="Acumulado",
            mode="lines+markers",line=dict(color="#E63946",width=2),marker=dict(size=6)))
        fig_acu.update_layout(height=360,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=11,color="#333"),margin=dict(l=10,r=10,t=38,b=10),
            xaxis=dict(showgrid=False,type="category"),yaxis=dict(gridcolor="#F0F0F0"),
            legend=dict(orientation="h",y=-0.25,x=0.5,xanchor="center"),
            title=dict(text="Saldo y Acumulado por Año",font=dict(size=13,color="#1B3A6B"),x=0))
        plot(fig_acu,use_container_width=True)
    st.dataframe(cos[["Año","Saldo Fmt","Facturas","% Part"]].rename(columns={"Saldo Fmt":"Saldo"}),
                 use_container_width=True,hide_index=True)
    if col_ar and col_ar in df.columns:
        section("📬 Por Año de Radicación")
        cos_r=df.groupby(col_ar).agg(Saldo=("Saldo Actual","sum"),Facturas=("Saldo Actual","count")).reset_index()
        cos_r.columns=["Año Rad","Saldo","Facturas"]; cos_r["Año Rad"]=cos_r["Año Rad"].astype(str)
        cos_r["Fmt"]=cos_r["Saldo"].apply(fmt); cos_r=cos_r.sort_values("Año Rad")
        fig_cr=go.Figure(go.Bar(x=cos_r["Año Rad"],y=cos_r["Saldo"],marker_color=PALETTE[:len(cos_r)],
            text=cos_r["Fmt"],textposition="outside"))
        fig_cr.update_layout(height=320,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=11,color="#333"),margin=dict(l=10,r=10,t=38,b=10),
            xaxis=dict(showgrid=False,type="category"),yaxis=dict(gridcolor="#F0F0F0"),
            title=dict(text="Saldo por Año de Radicación",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
        plot(fig_cr,use_container_width=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 12 — JURÍDICO
# ══════════════════════════════════════════════════════════════════════════════
def seccion_juridico(df):
    section("⚖️ Análisis Jurídico")
    if not tiene(df,"juridico_estado") and not tiene(df,"juridico_proceso"):
        st.info("No se encontraron columnas jurídicas."); return
    mask=pd.Series([False]*len(df),index=df.index)
    if tiene(df,"juridico_estado"):
        mask=mask|(~df[COLS_OPT["juridico_estado"]].astype(str).str.strip().str.lower().isin(["","nan","no aplica"]))
    if tiene(df,"juridico_proceso"):
        mask=mask|df[COLS_OPT["juridico_proceso"]].astype(str).str.strip().str.lower().isin(["si","sí","x","1","true"])
    dfj=df[mask].copy(); total=df["Saldo Actual"].sum(); tj=dfj["Saldo Actual"].sum(); pctj=(tj/total*100) if total else 0
    c1,c2,c3=st.columns(3)
    with c1: kpi_card("Total en Jurídico",fmt(tj),f"{len(dfj):,} registros","danger")
    with c2: kpi_card("% sobre Cartera",fmt_pct(pctj),"En litigio","danger")
    with c3:
        if tiene(df,"abogado"): kpi_card("Abogados Asignados",str(dfj[COLS_OPT["abogado"]].dropna().nunique()),"casos activos","primary")
    cl1,cl2=st.columns(2)
    with cl1:
        if tiene(df,"juridico_estado"):
            je=dfj.groupby(COLS_OPT["juridico_estado"])["Saldo Actual"].sum().reset_index()
            je.columns=["Estado","Saldo"]; je["Fmt"]=je["Saldo"].apply(fmt)
            fj=px.bar(je,x="Estado",y="Saldo",color="Estado",color_discrete_sequence=PALETTE,text="Fmt")
            fj.update_traces(textposition="outside")
            plot(fig_layout(fj,"Por Estado Jurídico"),use_container_width=True)
    with cl2:
        if tiene(df,"abogado"):
            ja=dfj.groupby(COLS_OPT["abogado"])["Saldo Actual"].sum().nlargest(10).reset_index()
            ja.columns=["Abogado","Saldo"]; ja["Fmt"]=ja["Saldo"].apply(fmt)
            fj2=px.bar(ja.sort_values("Saldo"),x="Saldo",y="Abogado",orientation="h",
                       color="Saldo",color_continuous_scale=["#A8DADC","#1B3A6B"],text="Fmt")
            fj2.update_traces(textposition="outside"); fj2.update_coloraxes(showscale=False)
            plot(fig_layout(fj2,"Por Abogado"),use_container_width=True)
    if not dfj.empty:
        cs=["Razon Social","Tipo De Empresa","Saldo Actual"]
        if tiene(df,"juridico_estado"): cs.append(COLS_OPT["juridico_estado"])
        if tiene(df,"abogado"):         cs.append(COLS_OPT["abogado"])
        if "# PROCESO" in dfj.columns:  cs.append("# PROCESO")
        det=dfj[[c for c in cs if c in dfj.columns]].copy()
        det["Saldo Actual"]=det["Saldo Actual"].apply(fmt)
        st.dataframe(det,use_container_width=True,hide_index=True)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 13 — PAGOS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_pagos(df, dp):
    section("💳 Análisis de Pagos Históricos")
    if dp is None: st.info("Carga el archivo de pagos para ver este análisis."); return
    tp=dp["Valor pagado"].sum()
    reg=dp[dp["Estado"].astype(str).str.strip().str.lower()=="registrado"]
    sinreg=dp[dp["Estado"].astype(str).str.strip().str.lower()=="sin registrar"]
    pct_ap=(reg["Valor pagado"].sum()/tp*100) if tp else 0
    c1,c2,c3,c4,c5=st.columns(5)
    with c1: kpi_card("Total Pagos",fmt(tp),f"{len(dp):,} registros","purple")
    with c2: kpi_card("Registrados",fmt(reg["Valor pagado"].sum()),f"{len(reg):,}","success")
    with c3: kpi_card("Recaudo por Aplicar",fmt(sinreg["Valor pagado"].sum()),f"{len(sinreg):,}","danger")
    with c4: kpi_card("% Aplicado",fmt_pct(pct_ap),"de lo recibido","success" if pct_ap>70 else "warning")
    with c5: kpi_card("Clientes c/Pago",str(dp["Razon Social"].nunique()),"clientes","primary")
    with st.expander("📊 Ver gráficos de pagos"):
        if dp["Fecha de pago"].notna().any():
            dp2=dp.copy(); dp2["Año"]=dp2["Fecha de pago"].dt.year
            dp2["Mes_Num"]=dp2["Fecha de pago"].dt.month; dp2["Mes_Nom"]=dp2["Fecha de pago"].dt.strftime("%b")
            def etiq(a):
                try: return "Acum. 2012-2023" if int(a)<=2023 else str(a)
                except: return str(a)
            dp2["Grupo"]=dp2["Año"].apply(etiq)
            g1=dp2.groupby(["Grupo","Estado"])["Valor pagado"].sum().reset_index()
            og=[g for g in ["Acum. 2012-2023","2024","2025","2026"] if g in g1["Grupo"].unique()]
            g1["_o"]=g1["Grupo"].apply(lambda x:og.index(x) if x in og else 99)
            g1=g1.sort_values("_o").drop(columns=["_o"]); g1["Fmt"]=g1["Valor pagado"].apply(fmt)
            fp1=px.bar(g1,x="Grupo",y="Valor pagado",color="Estado",barmode="group",
                       color_discrete_map={"Registrado":"#2D9B6F","Sin registrar":"#E63946"},
                       text="Fmt",category_orders={"Grupo":og})
            fp1.update_traces(textposition="outside")
            plot(fig_layout(fp1,"Pagos por período",340),use_container_width=True)
            d26=dp2[dp2["Año"].isin([2025,2026])].copy()
            if not d26.empty:
                MO=["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
                g2=d26.groupby(["Año","Mes_Num","Mes_Nom"])["Valor pagado"].sum().reset_index()
                g2=g2.sort_values(["Año","Mes_Num"]); g2["Año"]=g2["Año"].astype(str); g2["Fmt"]=g2["Valor pagado"].apply(fmt)
                fp2=px.bar(g2,x="Mes_Nom",y="Valor pagado",color="Año",barmode="group",
                           color_discrete_map={"2025":"#2E86AB","2026":"#E63946"},
                           text="Fmt",category_orders={"Mes_Nom":MO})
                fp2.update_traces(textposition="outside")
                plot(fig_layout(fp2,"Comparativo 2025 vs 2026",340),use_container_width=True)
    cl1,cl2=st.columns(2)
    with cl1:
        section("🏢 Top 10 Clientes Registrados")
        tr=reg.groupby("Razon Social")["Valor pagado"].sum().nlargest(10).reset_index()
        tr["Fmt"]=tr["Valor pagado"].apply(fmt)
        fr=go.Figure(go.Bar(x=tr["Valor pagado"],y=tr["Razon Social"],orientation="h",
            marker=dict(color=tr["Valor pagado"],colorscale=[[0,"#A8DADC"],[1,"#2D9B6F"]],showscale=False),
            text=tr["Fmt"],textposition="outside",textfont=dict(size=10)))
        fr.update_layout(height=420,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=160,t=38,b=10),
            xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
            yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(tr["Razon Social"].tolist()))),
            title=dict(text="Top 10 Clientes Registrados",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
        plot(fr,use_container_width=True)
    with cl2:
        section("⚠️ Top 10 Clientes por Aplicar")
        ts=sinreg.groupby("Razon Social")["Valor pagado"].sum().nlargest(10).reset_index()
        ts["Fmt"]=ts["Valor pagado"].apply(fmt)
        fs=go.Figure(go.Bar(x=ts["Valor pagado"],y=ts["Razon Social"],orientation="h",
            marker=dict(color=ts["Valor pagado"],colorscale=[[0,"#F4A261"],[1,"#E63946"]],showscale=False),
            text=ts["Fmt"],textposition="outside",textfont=dict(size=10)))
        fs.update_layout(height=420,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=10,color="#333"),margin=dict(l=10,r=160,t=38,b=10),
            xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
            yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(ts["Razon Social"].tolist()))),
            title=dict(text="Top 10 Clientes por Aplicar",font=dict(size=13,color="#1B3A6B"),x=0),showlegend=False)
        plot(fs,use_container_width=True)
    section("🔗 Cruce Pagos vs Cartera — Cartera Neta")
    if tiene(df,"nit") and "Nit" in dp.columns:
        sn=df.groupby(COLS_OPT["nit"])["Saldo Actual"].sum().reset_index(); sn.columns=["Nit","Saldo Cartera"]
        sinreg_nit=sinreg.groupby("Nit")["Valor pagado"].sum().reset_index(); sinreg_nit.columns=["Nit","Sin Registrar"]
        cr=pd.merge(sn,sinreg_nit,on="Nit",how="left").fillna(0)
        cr=pd.merge(cr,df[[COLS_OPT["nit"],"Razon Social"]].drop_duplicates().rename(columns={COLS_OPT["nit"]:"Nit"}),on="Nit",how="left")
        cr["Cartera Neta"]=cr["Saldo Cartera"]-cr["Sin Registrar"]
        cr["% Sin Reg/Cartera"]=(cr["Sin Registrar"]/cr["Saldo Cartera"]*100).replace([float("inf")],0).fillna(0).round(1)
        cr["Estado"]=cr["Cartera Neta"].apply(lambda v:"🔴 Pendiente" if v>0 else("🟢 Al día" if v==0 else "🟡 Sobrepago"))
        cr=cr.sort_values("Saldo Cartera",ascending=False)
        disp=cr[["Razon Social","Nit","Saldo Cartera","Sin Registrar","Cartera Neta","% Sin Reg/Cartera","Estado"]].copy()
        disp["Saldo Cartera"]=cr["Saldo Cartera"].apply(fmt); disp["Sin Registrar"]=cr["Sin Registrar"].apply(fmt); disp["Cartera Neta"]=cr["Cartera Neta"].apply(fmt)
        st.dataframe(disp.rename(columns={"Sin Registrar":"Pagos Sin Registrar"}),use_container_width=True,hide_index=True,height=400)
    section("📄 Detalle de Pagos")
    det=dp.copy(); det["Valor pagado"]=dp["Valor pagado"].apply(fmt)
    det["Fecha de pago"]=dp["Fecha de pago"].dt.strftime("%d/%m/%Y").fillna("")
    ef=st.selectbox("Filtrar estado:",["Todos","Registrado","Sin registrar"])
    if ef!="Todos": det=det[det["Estado"].astype(str).str.strip().str.lower()==ef.lower()]
    st.dataframe(det,use_container_width=True,hide_index=True,height=360)

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 14 — RESUMEN POR CLIENTE
# ══════════════════════════════════════════════════════════════════════════════
def seccion_resumen_cliente(df, dp):
    section("👤 Resumen Detallado por Cliente")
    clientes=["(Selecciona un cliente)"]+sorted(df["Razon Social"].unique().tolist())
    cliente=st.selectbox("🔍 Selecciona el cliente:",clientes,key="sel_cliente_resumen")
    if cliente=="(Selecciona un cliente)":
        st.info("👆 Selecciona un cliente para ver su resumen completo."); return
    dfc=df[df["Razon Social"]==cliente].copy()
    total=df["Saldo Actual"].sum(); s_cli=dfc["Saldo Actual"].sum(); pct=(s_cli/total*100) if total else 0
    n_fac=len(dfc); tipo=dfc["Tipo De Empresa"].mode()[0] if not dfc["Tipo De Empresa"].mode().empty else "N/A"
    nit=dfc[COLS_OPT["nit"]].iloc[0] if tiene(dfc,"nit") else "N/A"
    dp_cli=pd.DataFrame()
    if dp is not None and tiene(df,"nit") and "Nit" in dp.columns:
        nits=dfc[COLS_OPT["nit"]].astype(str).unique()
        dp_cli=dp[dp["Nit"].isin(nits)].copy()
    st.markdown(f"""<div style="background:linear-gradient(135deg,#1B3A6B,#2E86AB);border-radius:12px;
    padding:1rem 1.5rem;margin-bottom:1rem;color:white">
    <h3 style="margin:0;color:white;font-size:1.3rem">👤 {cliente}</h3>
    <p style="margin:.3rem 0 0;opacity:.85;font-size:.88rem">
    NIT: {nit} &nbsp;|&nbsp; Tipo: {tipo} &nbsp;|&nbsp;
    Saldo: <b>{fmt(s_cli)}</b> &nbsp;|&nbsp; Part.: <b>{pct:.1f}%</b> &nbsp;|&nbsp; Facturas: <b>{n_fac:,}</b>
    </p></div>""",unsafe_allow_html=True)

    cobro_cli=dfc[dfc["Estado de cartera"].astype(str).str.strip()=="Para cobro"]["Saldo Actual"].sum()
    pct_cobro_cli=(cobro_cli/s_cli*100) if s_cli else 0
    pct_cobro_tot=(df[df["Estado de cartera"].astype(str).str.strip()=="Para cobro"]["Saldo Actual"].sum()/total*100) if total else 0
    vencida_cli=dfc["_num_vencida"].sum() if "_num_vencida" in dfc.columns else 0
    pct_venc_cli=(vencida_cli/s_cli*100) if s_cli else 0
    pct_venc_tot=(df["_num_vencida"].sum()/total*100) if ("_num_vencida" in df.columns and total) else 0
    glosa_cli=dfc[COLS_OPT["glosa_saldo"]].sum() if tiene(dfc,"glosa_saldo") else 0
    pct_glo_cli=(glosa_cli/s_cli*100) if s_cli else 0
    pct_glo_tot=(df[COLS_OPT["glosa_saldo"]].sum()/total*100) if (tiene(df,"glosa_saldo") and total) else 0
    por_rad_cli=dfc[dfc["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()
    pct_rad_cli=(por_rad_cli/s_cli*100) if s_cli else 0
    prom_dias_cli=dfc["_dias_rad"].mean() if "_dias_rad" in dfc.columns else 0
    prom_dias_tot=df["_dias_rad"].mean() if "_dias_rad" in df.columns else 0

    k1,k2,k3,k4,k5=st.columns(5)
    with k1: kpi_card("Saldo Total",fmt(s_cli),f"{pct:.1f}% de la cartera","primary")
    with k2: kpi_card("Cartera en Cobro",fmt(cobro_cli),f"Cliente:{pct_cobro_cli:.1f}% | Total:{pct_cobro_tot:.1f}%","danger")
    with k3: kpi_card("Factura Vencida",fmt(vencida_cli),f"Cliente:{pct_venc_cli:.1f}% | Total:{pct_venc_tot:.1f}%","warning")
    with k4: kpi_card("Saldo Glosas",fmt(glosa_cli),f"Cliente:{pct_glo_cli:.1f}% | Total:{pct_glo_tot:.1f}%","warning")
    with k5: kpi_card("Por Radicar",fmt(por_rad_cli),fmt_pct(pct_rad_cli)+" del cliente","warning")

    if not dp_cli.empty:
        reg_cli=dp_cli[dp_cli["Estado"].astype(str).str.strip().str.lower()=="registrado"]["Valor pagado"].sum()
        sinreg_cli=dp_cli[dp_cli["Estado"].astype(str).str.strip().str.lower()=="sin registrar"]["Valor pagado"].sum()
        pct_ap_cli=(reg_cli/(reg_cli+sinreg_cli)*100) if (reg_cli+sinreg_cli) else 0
        k6,k7,k8,k9,k10=st.columns(5)
        with k6:  kpi_card("Pagos Registrados",fmt(reg_cli),f"{len(dp_cli[dp_cli['Estado'].str.lower()=='registrado']):,}","success")
        with k7:  kpi_card("Recaudo x Aplicar",fmt(sinreg_cli),f"{len(dp_cli[dp_cli['Estado'].str.lower()=='sin registrar']):,}","danger")
        with k8:  kpi_card("% Aplicado",fmt_pct(pct_ap_cli),"del total","success" if pct_ap_cli>70 else "warning")
        with k9:  kpi_card("Días Prom. Rad.",f"{prom_dias_cli:.0f} días" if prom_dias_cli else "N/A",f"IPS: {prom_dias_tot:.0f} días","warning" if prom_dias_cli>30 else "success")
        with k10: kpi_card("Max Factura",fmt(dfc["Saldo Actual"].max()),"mayor valor","default")

    sub_tabs=st.tabs(["📊 Cartera","🔍 Glosas","💳 Pagos","⏱️ Radicación","↩️ Devueltas","⚖️ Jurídico","📋 Detalle + Excel"])

    with sub_tabs[0]:
        c1,c2=st.columns(2)
        with c1:
            est_c=dfc.groupby("Estado de cartera")["Saldo Actual"].sum().reset_index()
            est_c["% Part"]=(est_c["Saldo Actual"]/s_cli*100).round(1); est_c["Fmt"]=est_c["Saldo Actual"].apply(fmt)
            est_c=est_c.sort_values("Saldo Actual",ascending=False)
            bc=[COLORS["danger"] if str(x).strip()=="Para cobro" else PALETTE[i%len(PALETTE)] for i,x in enumerate(est_c["Estado de cartera"])]
            fig_e=go.Figure(go.Bar(x=est_c["Saldo Actual"],y=est_c["Estado de cartera"],orientation="h",
                marker_color=bc,text=est_c["Fmt"],textposition="outside"))
            fig_e.update_layout(height=300,paper_bgcolor="white",plot_bgcolor="white",
                font=dict(family="Arial",size=10),margin=dict(l=10,r=140,t=38,b=10),
                xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),yaxis=dict(showgrid=False),
                title=dict(text="Por Estado",font=dict(size=12,color="#1B3A6B"),x=0),showlegend=False)
            plot(fig_e,use_container_width=True)
        with c2:
            ORDEN=["Por radicar","1-30","31-60","61-90","91-150","151-210","211-360","2 años","3 años","4 años","5 años","> 5 AÑOS"]
            int_c=dfc.groupby("Intervalo-Actual")["Saldo Actual"].sum().reset_index()
            int_c["_o"]=int_c["Intervalo-Actual"].apply(lambda v:next((i for i,o in enumerate(ORDEN) if str(v).strip().upper()==o.upper()),99))
            int_c=int_c.sort_values("_o").drop(columns=["_o"])
            int_c["% Part"]=(int_c["Saldo Actual"]/s_cli*100).round(1); int_c["Fmt"]=int_c["Saldo Actual"].apply(fmt)
            COLS_AG=["#2D9B6F","#57C98A","#F4D35E","#F4A261","#E76F51","#E63946","#9B2335","#6B0F1A","#4A0A12","#2D0508","#1A0005","#000000"]
            fig_i=go.Figure(go.Bar(x=int_c["Saldo Actual"],y=int_c["Intervalo-Actual"],orientation="h",
                marker_color=[COLS_AG[min(j,len(COLS_AG)-1)] for j in range(len(int_c))],
                text=int_c["Fmt"],textposition="outside"))
            fig_i.update_layout(height=300,paper_bgcolor="white",plot_bgcolor="white",
                font=dict(family="Arial",size=10),margin=dict(l=10,r=140,t=38,b=10),
                xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),
                yaxis=dict(showgrid=False,categoryorder="array",categoryarray=list(reversed(int_c["Intervalo-Actual"].tolist()))),
                title=dict(text="Por Intervalo",font=dict(size=12,color="#1B3A6B"),x=0),showlegend=False)
            plot(fig_i,use_container_width=True)
        st.markdown("**📊 Comparativo Cliente vs Cartera Total**")
        comp_df=pd.DataFrame({
            "Indicador":["% En Cobro","% Vencida","% Glosas","% Por Radicar"],
            "Cliente %":[pct_cobro_cli,pct_venc_cli,pct_glo_cli,pct_rad_cli],
            "Total %":[pct_cobro_tot,pct_venc_tot,pct_glo_tot,
                       df[df["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()/total*100 if total else 0],
        })
        fig_comp=go.Figure()
        fig_comp.add_trace(go.Bar(name="Cliente",x=comp_df["Indicador"],y=comp_df["Cliente %"],
            marker_color="#E63946",text=[f"{v:.1f}%" for v in comp_df["Cliente %"]],textposition="outside"))
        fig_comp.add_trace(go.Bar(name="Cartera Total",x=comp_df["Indicador"],y=comp_df["Total %"],
            marker_color="#2E86AB",text=[f"{v:.1f}%" for v in comp_df["Total %"]],textposition="outside"))
        fig_comp.update_layout(barmode="group",height=300,paper_bgcolor="white",plot_bgcolor="white",
            font=dict(family="Arial",size=11),margin=dict(l=10,r=10,t=38,b=10),
            xaxis=dict(showgrid=False),yaxis=dict(gridcolor="#F0F0F0",title="%"),
            legend=dict(orientation="h",y=-0.3,x=0.5,xanchor="center"),
            title=dict(text="Cliente vs Cartera Total (%)",font=dict(size=12,color="#1B3A6B"),x=0))
        plot(fig_comp,use_container_width=True)
        st.dataframe(est_c[["Estado de cartera","Fmt","% Part"]].rename(columns={"Estado de cartera":"Estado","Fmt":"Saldo"}),
                     use_container_width=True,hide_index=True)

    with sub_tabs[1]:
        cg=COLS_OPT.get("glosa_valor"); cgs=COLS_OPT.get("glosa_saldo"); col_ge=COLS_OPT.get("glosa_estado"); col_vf=COLS_OPT.get("vr_factura")
        has_g=(cg and cg in dfc.columns) or (cgs and cgs in dfc.columns)
        if not has_g: st.info("No se encontraron columnas de glosas.")
        else:
            dfc_g=dfc[(dfc[cg]>0) if (cg and cg in dfc.columns) else pd.Series(True,index=dfc.index)].copy()
            tg_c=dfc_g[cgs].sum() if (cgs and cgs in dfc_g.columns) else 0
            tv_c=dfc_g[cg].sum() if (cg and cg in dfc_g.columns) else 0
            idx_g=0
            if col_vf and col_vf in dfc.columns and dfc[col_vf].sum()>0:
                idx_g=(tv_c/dfc[col_vf].sum()*100)
            g1,g2,g3,g4=st.columns(4)
            with g1: kpi_card("Saldo Glosa s/Cartera",fmt(tg_c),fmt_pct(tg_c/s_cli*100 if s_cli else 0),"warning")
            with g2: kpi_card("Vr Glosa Reportada",fmt(tv_c),"Valor bruto","warning")
            with g3: kpi_card("Índice de Glosas",f"{idx_g:.2f}%","Glosa/Facturado","danger" if idx_g>10 else "warning")
            with g4: kpi_card("Registros c/Glosa",f"{len(dfc_g):,}","facturas afectadas","default")
            if col_ge and col_ge in dfc_g.columns and len(dfc_g)>0:
                gx=dfc_g.groupby(col_ge)[cgs if cgs and cgs in dfc_g.columns else cg].sum().reset_index()
                gx.columns=["Estado","Valor"]; gx["Fmt"]=gx["Valor"].apply(fmt)
                ce=["#2D9B6F" if any(x in str(e).lower() for x in ["conciliada","acta"]) else ("#F4A261" if "respuesta" in str(e).lower() else "#E63946") for e in gx["Estado"]]
                fgg=go.Figure(go.Bar(x=gx["Valor"],y=gx["Estado"],orientation="h",marker_color=ce,text=gx["Fmt"],textposition="outside"))
                fgg.update_layout(height=280,paper_bgcolor="white",plot_bgcolor="white",font=dict(family="Arial",size=10),
                    margin=dict(l=10,r=140,t=38,b=10),xaxis=dict(showgrid=True,gridcolor="#F0F0F0"),yaxis=dict(showgrid=False),
                    title=dict(text="Glosas por Estado",font=dict(size=12,color="#1B3A6B"),x=0),showlegend=False)
                plot(fgg,use_container_width=True)
            cols_gd=["Intervalo-Actual","Estado de cartera","Vr,Glosa Reportada","Saldo de Glosa según cartera","Estado De Glosa","Fecha_Recepcion","Dias de Glosa","Saldo Actual","OBSERVACION"]
            det_g=dfc_g[[c for c in cols_gd if c in dfc_g.columns]].copy()
            for nc in ["Vr,Glosa Reportada","Saldo de Glosa según cartera","Saldo Actual"]:
                if nc in det_g.columns:
                    det_g[nc]=pd.to_numeric(det_g[nc],errors="coerce").apply(lambda x:fmt(x) if pd.notna(x) else "")
            st.dataframe(det_g,use_container_width=True,hide_index=True,height=300)

    with sub_tabs[2]:
        if dp_cli.empty: st.info("No se encontraron pagos para este cliente.")
        else:
            reg_c=dp_cli[dp_cli["Estado"].astype(str).str.lower()=="registrado"]
            sin_c=dp_cli[dp_cli["Estado"].astype(str).str.lower()=="sin registrar"]
            tp_c=dp_cli["Valor pagado"].sum()
            pct_ap=(reg_c["Valor pagado"].sum()/tp_c*100) if tp_c else 0
            p1,p2,p3,p4=st.columns(4)
            with p1: kpi_card("Total Recibido",fmt(tp_c),f"{len(dp_cli):,}","purple")
            with p2: kpi_card("Registrados",fmt(reg_c["Valor pagado"].sum()),f"{len(reg_c):,}","success")
            with p3: kpi_card("Recaudo x Aplicar",fmt(sin_c["Valor pagado"].sum()),f"{len(sin_c):,}","danger")
            with p4: kpi_card("% Aplicado",fmt_pct(pct_ap),"del total","success" if pct_ap>70 else "warning")
            if dp_cli["Fecha de pago"].notna().any():
                dp_c2=dp_cli.copy(); dp_c2["Mes"]=dp_c2["Fecha de pago"].dt.to_period("M").astype(str)
                mens_c=dp_c2.groupby(["Mes","Estado"])["Valor pagado"].sum().reset_index().sort_values("Mes")
                fig_pc=px.bar(mens_c,x="Mes",y="Valor pagado",color="Estado",barmode="group",
                    color_discrete_map={"Registrado":"#2D9B6F","Sin registrar":"#E63946"},
                    text=mens_c["Valor pagado"].apply(fmt))
                fig_pc.update_traces(textposition="outside")
                plot(fig_layout(fig_pc,"Evolución Mensual de Pagos",340),use_container_width=True)
            dp_show=dp_cli.copy(); dp_show["Valor pagado"]=dp_cli["Valor pagado"].apply(fmt)
            dp_show["Fecha de pago"]=dp_cli["Fecha de pago"].dt.strftime("%d/%m/%Y").fillna("")
            st.dataframe(dp_show,use_container_width=True,hide_index=True,height=300)

    with sub_tabs[3]:
        if "_dias_rad" not in dfc.columns: st.info("No se encontraron fechas de radicación.")
        else:
            dr=dfc[dfc["_dias_rad"].notna()&(dfc["_dias_rad"]>=0)].copy()
            if dr.empty: st.info("Sin datos de radicación para este cliente.")
            else:
                prom_dr=dr["_dias_rad"].mean(); med_dr=dr["_dias_rad"].median(); max_dr=dr["_dias_rad"].max(); min_dr=dr["_dias_rad"].min()
                r1,r2,r3,r4=st.columns(4)
                with r1: kpi_card("Días Prom.",f"{prom_dr:.0f} días",f"IPS: {prom_dias_tot:.0f} días","warning" if prom_dr>30 else "success")
                with r2: kpi_card("Mediana",f"{med_dr:.0f} días","","default")
                with r3: kpi_card("Máximo",f"{max_dr:.0f} días","","danger" if max_dr>90 else "warning")
                with r4: kpi_card("Mínimo",f"{min_dr:.0f} días","","success")
                bins=[0,15,30,60,90,180,365,9999]; lbls=["0-15","16-30","31-60","61-90","91-180","181-365","+365"]
                dr["_rr"]=pd.cut(dr["_dias_rad"],bins=bins,labels=lbls,right=True)
                rng=dr.groupby("_rr",observed=True)["Saldo Actual"].agg(["sum","count"]).reset_index()
                rng.columns=["Rango","Saldo","Facturas"]; rng["Fmt"]=rng["Saldo"].apply(fmt)
                COLS_R=["#2D9B6F","#57C98A","#F4D35E","#F4A261","#E76F51","#E63946","#7B1D1D"]
                fig_rr=go.Figure(go.Bar(x=rng["Rango"],y=rng["Saldo"],marker_color=COLS_R[:len(rng)],
                    text=rng["Fmt"],textposition="outside"))
                plot(fig_layout(fig_rr,"Saldo por Rango de Días",300),use_container_width=True)

    with sub_tabs[4]:
        col_dev=COLS_OPT.get("devuelta"); col_devj=COLS_OPT.get("dev_juridica"); col_mot=COLS_OPT.get("motivo_dev")
        mask_d=pd.Series([False]*len(dfc),index=dfc.index)
        if col_dev  and col_dev  in dfc.columns: mask_d=mask_d|dfc[col_dev].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])
        if col_devj and col_devj in dfc.columns: mask_d=mask_d|dfc[col_devj].astype(str).str.strip().str.upper().isin(["SI","SÍ","S","X","1","TRUE"])
        dfd_c=dfc[mask_d].copy()
        if dfd_c.empty: st.success("✅ No se registran facturas devueltas para este cliente.")
        else:
            t_d=dfd_c["Saldo Actual"].sum()
            d1,d2,d3=st.columns(3)
            with d1: kpi_card("Total Devuelto",fmt(t_d),fmt_pct(t_d/s_cli*100 if s_cli else 0),"danger")
            with d2: kpi_card("Facturas Devueltas",f"{len(dfd_c):,}","registros","warning")
            with d3: kpi_card("Promedio",fmt(t_d/len(dfd_c) if len(dfd_c) else 0),"por factura","default")
            cols_dv=["Intervalo-Actual","Estado de cartera","Saldo Actual"]
            if col_mot  and col_mot  in dfd_c.columns: cols_dv.append(col_mot)
            if "OBSERVACION" in dfd_c.columns: cols_dv.append("OBSERVACION")
            det_dv=dfd_c[[c for c in cols_dv if c in dfd_c.columns]].copy()
            det_dv["Saldo Actual"]=det_dv["Saldo Actual"].apply(fmt)
            st.dataframe(det_dv,use_container_width=True,hide_index=True,height=320)

    with sub_tabs[5]:
        col_je=COLS_OPT.get("juridico_estado"); col_jp=COLS_OPT.get("juridico_proceso"); col_ab=COLS_OPT.get("abogado")
        mask_j=pd.Series([False]*len(dfc),index=dfc.index)
        if col_je and col_je in dfc.columns: mask_j=mask_j|(~dfc[col_je].astype(str).str.strip().str.lower().isin(["","nan","no aplica"]))
        if col_jp and col_jp in dfc.columns: mask_j=mask_j|dfc[col_jp].astype(str).str.strip().str.lower().isin(["si","sí","x","1","true"])
        dfj_c=dfc[mask_j].copy()
        if dfj_c.empty: st.success("✅ No se registran procesos jurídicos activos para este cliente.")
        else:
            tj_c=dfj_c["Saldo Actual"].sum()
            j1,j2,j3=st.columns(3)
            with j1: kpi_card("Total en Jurídico",fmt(tj_c),fmt_pct(tj_c/s_cli*100 if s_cli else 0),"danger")
            with j2: kpi_card("Procesos Activos",f"{len(dfj_c):,}","registros","danger")
            with j3:
                if col_ab and col_ab in dfj_c.columns: kpi_card("Abogados",str(dfj_c[col_ab].dropna().nunique()),"asignados","primary")
            csj=["Estado de cartera","Saldo Actual","Intervalo-Actual"]
            if col_je and col_je in dfj_c.columns: csj.append(col_je)
            if col_ab and col_ab in dfj_c.columns: csj.append(col_ab)
            if "# PROCESO" in dfj_c.columns: csj.append("# PROCESO")
            if "OBSERVACION" in dfj_c.columns: csj.append("OBSERVACION")
            det_j=dfj_c[[c for c in csj if c in dfj_c.columns]].copy()
            det_j["Saldo Actual"]=det_j["Saldo Actual"].apply(fmt)
            st.dataframe(det_j,use_container_width=True,hide_index=True,height=320)

    with sub_tabs[6]:
        st.markdown(f"**📄 Detalle completo — {cliente}**")
        cols_exp=["Intervalo-Actual","Estado de cartera","Saldo Actual","Tipo De Empresa",
                  "Vr factura","Vr,Glosa Reportada","Saldo de Glosa según cartera",
                  "Estado De Glosa","Factura vencida","Por vencer","Dias Mora",
                  "Fecha fac","Fecha rad","Estado juridico","ABOGADO",
                  "Motivo de devolucion","OBSERVACION","Observacion de la factura"]
        det_exp=dfc[[c for c in cols_exp if c in dfc.columns]].copy()
        for nc in ["Saldo Actual","Vr factura","Vr,Glosa Reportada","Saldo de Glosa según cartera"]:
            if nc in det_exp.columns:
                det_exp[nc]=pd.to_numeric(det_exp[nc],errors="coerce").apply(lambda x:fmt(x) if pd.notna(x) else "")
        for dc_ in ["Fecha fac","Fecha rad"]:
            if dc_ in det_exp.columns:
                det_exp[dc_]=pd.to_datetime(det_exp[dc_],errors="coerce").dt.strftime("%d/%m/%Y")
        st.dataframe(det_exp.sort_values("Saldo Actual",ascending=False),use_container_width=True,hide_index=True,height=360)
        st.markdown("---")
        section("📥 Exportar en Excel")
        import io
        buffer=io.BytesIO()
        nombre=f"Resumen_{cliente[:30].replace(' ','_').replace('/','_')}.xlsx"
        with pd.ExcelWriter(buffer,engine="openpyxl") as writer:
            pd.DataFrame({
                "Indicador":["Cliente","NIT","Tipo","Saldo Total","Part. %","En Cobro","% Cobro","Vencida","Glosas","Por Radicar","Días Prom. Rad."],
                "Valor":[cliente,nit,tipo,s_cli,round(pct,2),cobro_cli,round(pct_cobro_cli,2),vencida_cli,glosa_cli,por_rad_cli,round(prom_dias_cli,0) if prom_dias_cli else 0]
            }).to_excel(writer,sheet_name="KPIs",index=False)
            dfc.groupby("Estado de cartera")["Saldo Actual"].sum().reset_index().to_excel(writer,sheet_name="Por Estado",index=False)
            dfc.groupby("Intervalo-Actual")["Saldo Actual"].sum().reset_index().to_excel(writer,sheet_name="Por Intervalo",index=False)
            if tiene(dfc,"glosa_saldo"):
                dfc_g_exp=dfc[dfc[COLS_OPT.get("glosa_valor","")]>0] if tiene(dfc,"glosa_valor") else dfc
                gcols=["Intervalo-Actual","Estado de cartera","Vr,Glosa Reportada","Saldo de Glosa según cartera","Estado De Glosa","Fecha_Recepcion","Saldo Actual"]
                dfc_g_exp[[c for c in gcols if c in dfc_g_exp.columns]].to_excel(writer,sheet_name="Glosas",index=False)
            if not dp_cli.empty:
                dp_e=dp_cli.copy(); dp_e["Fecha de pago"]=dp_e["Fecha de pago"].dt.strftime("%d/%m/%Y")
                dp_e.to_excel(writer,sheet_name="Pagos",index=False)
            dfc_raw=dfc[[c for c in cols_exp if c in dfc.columns]].copy()
            for dc_ in ["Fecha fac","Fecha rad"]:
                if dc_ in dfc_raw.columns: dfc_raw[dc_]=pd.to_datetime(dfc_raw[dc_],errors="coerce").dt.strftime("%d/%m/%Y")
            dfc_raw.to_excel(writer,sheet_name="Detalle Facturas",index=False)
        buffer.seek(0)
        st.download_button(label=f"⬇️ Descargar Excel — {cliente[:25]}",data=buffer,file_name=nombre,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",type="primary")

# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 15 — HALLAZGOS
# ══════════════════════════════════════════════════════════════════════════════
def seccion_hallazgos(df):
    total=df["Saldo Actual"].sum()
    cobro=df[df["Estado de cartera"].astype(str).str.strip()=="Para cobro"]
    pct_c=(cobro["Saldo Actual"].sum()/total*100) if total else 0
    top10p=(df.groupby("Razon Social")["Saldo Actual"].sum().nlargest(10).sum()/total*100) if total else 0
    n=df["Razon Social"].nunique()
    section("🧠 Hallazgos Clave de la Cartera")
    col_h1,col_h2=st.columns(2)
    with col_h1:
        if top10p>60:   insight(f"Alta concentración: Top 10 = {top10p:.1f}% ({n} clientes). Riesgo sistémico de liquidez.","danger")
        elif top10p>40: insight(f"Concentración moderada: Top 10 = {top10p:.1f}%. Monitoreo continuo.","warning")
        else:           insight(f"Cartera diversificada: Top 10 = {top10p:.1f}%.","success")
        if pct_c>30:    insight(f"Nivel crítico 'Para cobro': {pct_c:.1f}% ({fmt(cobro['Saldo Actual'].sum())}). Acción inmediata.","danger")
        elif pct_c>15:  insight(f"Cartera en cobro en alerta: {pct_c:.1f}%.","warning")
        else:           insight(f"Cartera en cobro controlada: {pct_c:.1f}%.","success")
        if tiene(df,"glosa_saldo"):
            tg=df[COLS_OPT["glosa_saldo"]].sum(); pg=(tg/total*100) if total else 0
            if pg>10: insight(f"Glosas = {pg:.1f}% ({fmt(tg)}). Impacto crítico en flujo de caja.","danger")
            elif pg>0:insight(f"Saldo glosas: {fmt(tg)} ({pg:.1f}%). Gestión recomendada.","warning")
    with col_h2:
        liq=df[df["Estado de cartera"].astype(str).str.lower().str.contains("liquidaci",na=False)]
        if len(liq)>0: insight(f"{liq['Razon Social'].nunique()} clientes en liquidación: {fmt(liq['Saldo Actual'].sum())}. Activar acreencias.","danger")
        jur=df[df["Estado de cartera"].astype(str).str.strip()=="En proceso juridico"]
        if len(jur)>0: insight(f"Cartera jurídica: {fmt(jur['Saldo Actual'].sum())} ({jur['Razon Social'].nunique()} clientes).","warning")
        pr=df[df["Intervalo-Actual"].astype(str).str.strip().str.lower()=="por radicar"]["Saldo Actual"].sum()
        if pr>0: insight(f"Por radicar: {fmt(pr)} ({fmt_pct(pr/total*100 if total else 0)}). Flujo detenido.","warning")
        v_crit=df[df["Intervalo-Actual"].astype(str).str.strip().isin(["2 años","3 años","4 años","5 años","> 5 AÑOS","211-360"])]["Saldo Actual"].sum()
        if v_crit>0: insight(f"Cartera +211 días: {fmt(v_crit)} ({fmt_pct(v_crit/total*100 if total else 0)}). Evaluar provisión NIIF.","danger")
    st.markdown("---")
    section("💡 Recomendaciones Estratégicas")
    recs=[]
    if pct_c>15: recs.append(("danger","Priorizar gestión de cobro",
        f"Activar cobro sobre {len(cobro):,} registros 'Para cobro' ({fmt(cobro['Saldo Actual'].sum())}). Gestores por EPS."))
    if top10p>50:
        t1=df.groupby("Razon Social")["Saldo Actual"].sum().idxmax()
        recs.append(("warning","Gestión focalizada Top deudores",f"Negociar acuerdos escalonados. Mayor exposición: {t1}."))
    if tiene(df,"juridico_estado"):
        recs.append(("warning","Revisión jurídica mensual","Informe de avance por apoderado. Evaluar acuerdos ante SuperSalud."))
    if tiene(df,"glosa_saldo") and df[COLS_OPT["glosa_saldo"]].sum()>0:
        recs.append(("warning","Gestión de glosas","Cronograma quincenal con auditoría médica. Priorizar por valor y fecha."))
    recs.append(("info","Comité mensual de cartera","KPIs: cobro, aging, concentración, jurídico, glosas, pagos. Informe trimestral a junta."))
    recs.append(("success","Acuerdos SuperSalud/ADRES","Conciliación extrajudicial y compensaciones con ADRES para EPS con cartera crítica."))
    for kind,titulo,texto in recs:
        st.markdown(f'<div class="insight-box {kind if kind!="info" else ""}"><b>{titulo}:</b> {texto}</div>',unsafe_allow_html=True)

    st.markdown("---")
    section("🔎 Ficha de Análisis por Cliente")
    clientes_lista=["(Selecciona un cliente)"]+sorted(df["Razon Social"].unique().tolist())
    cliente_sel=st.selectbox("Selecciona el cliente:",clientes_lista,key="hallazgos_cliente")
    if cliente_sel=="(Selecciona un cliente)":
        st.info("👆 Selecciona un cliente para ver su análisis individual."); return
    dfc=df[df["Razon Social"]==cliente_sel].copy()
    s_cli=dfc["Saldo Actual"].sum(); pct_cli=(s_cli/total*100) if total else 0
    pct_cobro_cli=(dfc[dfc["Estado de cartera"].astype(str).str.strip()=="Para cobro"]["Saldo Actual"].sum()/s_cli*100) if s_cli else 0
    v_crit_cli=dfc[dfc["Intervalo-Actual"].astype(str).str.strip().isin(["2 años","3 años","4 años","5 años","> 5 AÑOS","211-360"])]["Saldo Actual"].sum()
    pct_crit_cli=(v_crit_cli/s_cli*100) if s_cli else 0
    riesgo="ALTO 🔴" if (pct_cobro_cli>40 or pct_crit_cli>30) else ("MEDIO 🟡" if (pct_cobro_cli>15 or pct_crit_cli>15) else "BAJO 🟢")
    analisis_box(f"<b>{cliente_sel}</b> — Saldo: <b>{fmt(s_cli)}</b> ({pct_cli:.1f}% del total) | "
                 f"{len(dfc):,} facturas | Estado ppal: <b>{dfc['Estado de cartera'].mode()[0] if not dfc['Estado de cartera'].mode().empty else 'N/A'}</b> | "
                 f"Intervalo ppal: <b>{dfc['Intervalo-Actual'].mode()[0] if not dfc['Intervalo-Actual'].mode().empty else 'N/A'}</b><br><br>"
                 f"En cobro: <b>{pct_cobro_cli:.1f}%</b> | Antigüedad crítica: <b>{pct_crit_cli:.1f}%</b> | "
                 f"<b>Nivel de riesgo: {riesgo}</b>")
    section(f"⚠️ Alertas — {cliente_sel}")
    al_cli=[]
    if pct_cobro_cli>30: al_cli.append(("danger",f"{pct_cobro_cli:.1f}% en 'Para cobro' ({fmt(dfc[dfc['Estado de cartera'].astype(str).str.strip()=='Para cobro']['Saldo Actual'].sum())}). Gestión urgente."))
    if pct_crit_cli>20:  al_cli.append(("danger",f"Cartera +211 días: {fmt(v_crit_cli)} ({pct_crit_cli:.1f}%). Alta probabilidad de irrecuperabilidad."))
    if tiene(df,"glosa_saldo") and dfc[COLS_OPT["glosa_saldo"]].sum()>0:
        al_cli.append(("warning",f"Saldo de glosas: {fmt(dfc[COLS_OPT['glosa_saldo']].sum())}."))
    if not al_cli: al_cli.append(("success","No se identificaron alertas críticas. Cartera bajo control."))
    for kind,txt in al_cli:
        st.markdown(f'<div class="insight-box {kind}">{"🔴" if kind=="danger" else ("🟡" if kind=="warning" else "🟢")} {txt}</div>',unsafe_allow_html=True)
    section(f"💡 Recomendaciones — {cliente_sel}")
    rc=[]
    if pct_cobro_cli>15: rc.append(("danger","Activar cobro inmediato",f"Contactar tesorería de {cliente_sel}. Proponer acuerdo de pago a 30-60-90 días."))
    if pct_crit_cli>15:  rc.append(("warning","Evaluar provisión contable",f"Cartera +211 días ({fmt(v_crit_cli)}) con alta probabilidad de pérdida. Presentar al área contable."))
    if tiene(df,"glosa_saldo") and dfc[COLS_OPT["glosa_saldo"]].sum()>0:
        rc.append(("warning","Gestión de glosas",f"Coordinar con auditoría médica respuesta a glosas. Solicitar reunión de conciliación con {cliente_sel}."))
    if not rc: rc.append(("success","Mantener seguimiento","Cliente con indicadores controlados. Comunicación periódica y radicación oportuna."))
    for kind,titulo,texto in rc:
        st.markdown(f'<div class="insight-box {kind if kind!="info" else ""}"><b>{titulo}:</b> {texto}</div>',unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    _kc["n"] = 0
    st.markdown("""<div class="main-header">
        <h1>🏥 Tablero de Cartera — IPS</h1>
        <p>Análisis integral de cartera · Comité de Cartera y Junta Directiva</p>
    </div>""", unsafe_allow_html=True)

    cu1,cu2=st.columns(2)
    with cu1: f_cartera=st.file_uploader("📂 Archivo de Cartera (.xlsx)",type=["xlsx","xls"])
    with cu2: f_pagos=st.file_uploader("💳 Archivo de Pagos Históricos (.xlsx)",type=["xlsx","xls"],
                                        help="Columnas: Nit, Razon Social, Fecha de pago, Valor pagado, Estado")
    if f_cartera is None:
        st.info("👆 Carga el archivo de cartera para iniciar el análisis."); return

    with st.spinner("Procesando cartera..."):
        df_raw,err=cargar_cartera(f_cartera)
    if err: st.error(f"❌ {err}"); return

    dp=None
    if f_pagos is not None:
        with st.spinner("Procesando pagos..."):
            dp,err_p=cargar_pagos(f_pagos)
        if err_p: st.warning(f"⚠️ Pagos: {err_p}")

    if "tab_idx" not in st.session_state: st.session_state.tab_idx=0
    TAB_KEYS=["resumen","kpis","aging","tipo","glosas","indice_glosas",
              "radicacion","devueltas","acuerdos","calor","cosecha",
              "juridico","pagos","resumen_cliente","hallazgos"]
    tab_activa=TAB_KEYS[min(st.session_state.get("tab_idx",0),len(TAB_KEYS)-1)]

    filtros=build_sidebar(df_raw,tab_activa)
    df=aplicar_filtros(df_raw,filtros)

    # Botón exportar Word en sidebar
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📄 Exportar para NotebookLM")
    if len(df) < len(df_raw):
        st.sidebar.info(f"⚠️ Se exportarán **{len(df):,}** registros (filtros activos)")
    else:
        st.sidebar.caption(f"Se exportarán todos los registros: {len(df_raw):,}")
    if st.sidebar.button("⬇️ Generar Informe Word",type="primary",use_container_width=True):
        with st.spinner("Generando informe Word..."):
            buf_w,err_w=_generar_word_interno(df,dp)  # df = datos filtrados
        if err_w:
            st.sidebar.error(f"❌ {err_w}\n\nInstala con:\npip install python-docx")
        else:
            st.sidebar.download_button(
                label="📥 Descargar (.docx)",
                data=buf_w,
                file_name=f"Informe_Cartera_{pd.Timestamp.now().strftime('%Y%m%d')}{'_FILTRADO' if len(df)<len(df_raw) else ''}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    st.sidebar.caption("14 secciones · KPIs + tablas + hallazgos\nCompatible con NotebookLM")

    if df.empty: st.warning("⚠️ Sin datos con filtros activos."); return
    if len(df)<len(df_raw): st.info(f"📊 Mostrando **{len(df):,}** de **{len(df_raw):,}** registros")

    TABS=["📋 Resumen Ejecutivo","📈 KPIs","📊 Envejecimiento y Vencimientos",
          "🏢 Tipo Empresa","🔍 Glosas","📊 Índice Glosas","⏱️ Radicación",
          "↩️ Devueltas","🤝 Acuerdos y PGP","🗺️ Mapa de Calor",
          "📅 Cosecha","⚖️ Jurídico","💳 Pagos","👤 Resumen x Cliente","🧠 Hallazgos"]
    tabs=st.tabs(TABS)

    with tabs[0]:  st.session_state.tab_idx=0;  total,pct_cobro,pct_vencida,conc_top10=seccion_resumen(df,dp)
    with tabs[1]:  st.session_state.tab_idx=1;  seccion_kpis(df,dp,total,pct_cobro,pct_vencida,conc_top10)
    with tabs[2]:  st.session_state.tab_idx=2;  seccion_aging_vencimientos(df)
    with tabs[3]:  st.session_state.tab_idx=3;  seccion_tipo_empresa(df)
    with tabs[4]:  st.session_state.tab_idx=4;  seccion_glosas(df)
    with tabs[5]:  st.session_state.tab_idx=5;  seccion_indice_glosas(df)
    with tabs[6]:  st.session_state.tab_idx=6;  seccion_radicacion(df)
    with tabs[7]:  st.session_state.tab_idx=7;  seccion_devueltas(df)
    with tabs[8]:  st.session_state.tab_idx=8;  seccion_acuerdos_pgp(df)
    with tabs[9]:  st.session_state.tab_idx=9;  seccion_mapa_calor(df)
    with tabs[10]: st.session_state.tab_idx=10; seccion_cosecha(df)
    with tabs[11]: st.session_state.tab_idx=11; seccion_juridico(df)
    with tabs[12]: st.session_state.tab_idx=12; seccion_pagos(df,dp)
    with tabs[13]: st.session_state.tab_idx=13; seccion_resumen_cliente(df,dp)
    with tabs[14]: st.session_state.tab_idx=14; seccion_hallazgos(df)

    st.markdown("---")
    st.markdown("<div style='text-align:center;color:#6C757D;font-size:.8rem'>"
                "Tablero de Cartera IPS · Análisis Financiero · Uso interno gerencial</div>",
                unsafe_allow_html=True)

if __name__=="__main__":
    main()